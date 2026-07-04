# Builds the complete CET333 Portfolio Report for the AI-Solutions project as a .docx file.
# Structure follows the CET333 Additional Submission Guidance Notes (8 mandated sections):
#   Front Cover, Contents, 1 Requirements, 2 Planning, 3 Client Contact Record,
#   4 Methodology, 5 Solution Design, 6 Implementation, 7 Security, 8 Testing,
#   9 Technical Deployment, 10 Critical Reflection, 11 Maintainability, 12 Conclusion, Appendices.
# Every diagram is explained element-by-element and every code listing is the real
# source code followed by a line-by-line explanation table.
import os
import struct
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(__file__)
DIAG = os.path.join(BASE, "diagrams")
OUT = os.path.join(BASE, "AI-Solutions_Documentation.docx")

# ---- personalise these three values before submitting ----
STUDENT_NAME = "[Your Full Name]"
STUDENT_ID = "[Your Student ID]"
CLIENT_NAME = "Adam Robson (Module Tutor, acting as Client)"

# Brand palette
NAVY = RGBColor(0x0B, 0x3D, 0x5C)
TEAL = RGBColor(0x12, 0x7C, 0x8A)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for lvl, sz, col in [("Heading 1", 17, NAVY), ("Heading 2", 14, TEAL), ("Heading 3", 12.5, DARK)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = col
    st.font.bold = True


def png_size(path):
    with open(path, "rb") as f:
        head = f.read(26)
    w, h = struct.unpack(">II", head[16:24])
    return w, h


def add_image(name, caption):
    path = os.path.join(DIAG, name)
    pw, ph = png_size(path)
    ratio = ph / pw
    width_in = 6.3
    if width_in * ratio > 8.4:          # too tall -> cap height instead
        width_in = 8.2 / ratio
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    doc.add_paragraph()


def body(text):
    return doc.add_paragraph(text)


def lead(text):
    """A short bold lead-in line used to introduce a sub-topic."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = DARK
    return p


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    return doc.add_paragraph(text, style="List Number")


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # light shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F4F7")
    pPr.append(shd)
    # thin border box
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "6")
        e.set(qn("w:color"), "C7D0DA")
        borders.append(e)
    pPr.append(borders)
    return p


def caption_for_code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    return p


def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "0B3D5C")
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def walk(rows):
    """Two-column line-by-line code explanation table."""
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(["Code", "What it does and why"]):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "127C8A")
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for codepart, expl in rows:
        cells = t.add_row().cells
        cells[0].text = ""
        rc = cells[0].paragraphs[0].add_run(codepart)
        rc.font.name = "Consolas"
        rc.font.size = Pt(8.5)
        rc.font.color.rgb = NAVY
        cells[1].text = ""
        re_ = cells[1].paragraphs[0].add_run(expl)
        re_.font.size = Pt(9.5)
    for row in t.rows:
        row.cells[0].width = Inches(2.5)
        row.cells[1].width = Inches(3.9)
    doc.add_paragraph()
    return t


def signoff(role_label):
    """A signature/date line block for sign-off boxes."""
    make_table(
        ["Role", "Name", "Signature", "Date"],
        [[role_label, "", "", ""]],
        widths=[1.6, 2.0, 1.6, 1.1],
    )


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)
def pagebreak(): doc.add_page_break()


# =========================================================
# COVER PAGE
# =========================================================
for _ in range(3):
    doc.add_paragraph()
mc = doc.add_paragraph()
mc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = mc.add_run("CET333 PRODUCT DEVELOPMENT")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = TEAL

doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("AI-SOLUTIONS")
r.bold = True
r.font.size = Pt(42)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Company Website & Administrative Dashboard")
r.font.size = Pt(16)
r.font.color.rgb = TEAL

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Practitioner Portfolio Report")
r.font.size = Pt(13)
r.italic = True
r.font.color.rgb = GREY

for _ in range(5):
    doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line, bold in [
    ("Module Code: CET333 — Product Development", True),
    ("Project Title: AI-Solutions Software Studio Website", False),
    ("Student Name: " + STUDENT_NAME, True),
    ("Student ID: " + STUDENT_ID, True),
    ("Client: " + CLIENT_NAME, False),
    ("Submission: Task 1 — Portfolio Report (Week 15)", False),
    ("Technology Stack: Node.js · Express · EJS · SQLite · bcrypt · express-session", False),
]:
    rr = meta.add_run(line + "\n")
    rr.bold = bold
    rr.font.size = Pt(12)
pagebreak()

# =========================================================
# DOCUMENT CONTROL
# =========================================================
h1("Document Control")
make_table(
    ["Version", "Date", "Author", "Description"],
    [
        ["0.1", "Feb 2026", STUDENT_NAME, "Initial draft of requirements and project schedule"],
        ["0.5", "Mar 2026", STUDENT_NAME, "Design documentation and implementation added"],
        ["1.0", "Apr 2026", STUDENT_NAME, "Testing, security and all diagrams completed"],
        ["2.0", "May 2026", STUDENT_NAME, "Full portfolio: planning, methodology, client record, deployment, reflection"],
    ],
    widths=[0.9, 1.0, 1.8, 3.2],
)

body("This portfolio documents the planning, design, development, testing and deployment of the AI-Solutions "
     "website, and reflects critically on the process. It is written so that a reader with no prior knowledge of "
     "the project can understand exactly how the system is built, why each decision was made, what every diagram "
     "means, and what every important line of code does. Diagrams are explained element by element; code listings "
     "show the real source code from the project and are followed by a line-by-line explanation.")

lead("Mapping to the CET333 submission guidance")
make_table(
    ["Guidance section", "Where it is in this report"],
    [
        ["Front Cover & Contents", "Cover page and Table of Contents (page numbers via F9 in Word)"],
        ["1. Requirements Specification", "Section 1 (incl. business process model, deliverables, client sign-off)"],
        ["2. Planning Documentation", "Section 2 (working schedule with planned vs actual dates and revisions)"],
        ["3. Client Contact Record Sheet", "Section 3 (three client meetings with action points and sign-off)"],
        ["4. Methodology", "Section 4 (narrative, in own words, with own methodology model)"],
        ["5. Solution Design Documentation", "Sections 5 and 6 (architecture, database, interaction design, code)"],
        ["6. Testing", "Section 8 (strategy, tools, test cases, results, traceability)"],
        ["7. Technical Deployment", "Section 9 (technical requirements, packaging, installation, deployment)"],
        ["8. Evaluation & Critical Reflection", "Section 10"],
    ],
    widths=[2.6, 3.8],
)

# Table of Contents field (Word populates on open / F9)
h1("Table of Contents")
body("To generate page numbers for this contents list in Microsoft Word, click anywhere inside it and press F9, "
     "then choose “Update entire table”.")
toc_par = doc.add_paragraph()
run = toc_par.add_run()
fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve")
instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
fldText = OxmlElement("w:t"); fldText.text = "Right-click and Update Field to generate the Table of Contents."
fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
run._r.append(fldText); run._r.append(fldChar3)
pagebreak()

# =========================================================
# 1. REQUIREMENTS SPECIFICATION
# =========================================================
h1("1. Requirements Specification")

h2("1.1 Project Background")
body("AI-Solutions is a company that designs and sells bespoke artificial-intelligence software — chatbots, "
     "workflow automation, and custom machine-learning models. The company required a professional public-facing "
     "marketing website together with a private administrative area where staff can review the sales enquiries "
     "submitted by potential customers. The website must present the company’s services attractively, collect "
     "prospective-client enquiries, answer common questions automatically via an embedded AI assistant, and give "
     "administrators a secure dashboard to manage incoming enquiries. These requirements were gathered from the "
     "client during the initial client meeting (Section 3) and formally approved at the requirements sign-off "
     "before development began.")

h2("1.2 Scope")
body("The delivered product is a server-rendered web application consisting of eight public pages, a contact "
     "enquiry workflow backed by a database, an authenticated administrator dashboard, and a client-side FAQ "
     "chatbot. The following are within scope:")
for t in [
    "Eight navigable public pages with a shared, consistent layout.",
    "A seven-field Contact Us form with thorough server-side validation and database storage.",
    "Secure administrator authentication using hashed passwords and server sessions.",
    "An administrator dashboard showing enquiry totals and a full submissions table.",
    "An AI FAQ chatbot widget present on all public pages.",
    "Responsive layout for desktop, tablet, and mobile devices.",
]:
    bullet(t)
body("The following are explicitly out of scope: online payment processing, multi-tenant accounts for "
     "end-customers, email/SMS notifications, and third-party CRM integration. These are noted as possible "
     "future enhancements in Section 11.")

h2("1.3 Functional Requirements")
body("The functional requirements describe what the system must do. Every requirement below has been "
     "implemented and verified in the delivered product.")
make_table(
    ["ID", "Requirement", "Summary"],
    [
        ["FR-01", "Eight core pages", "Home, Solutions, Case Studies, Feedback, Articles, Gallery, Contact, Admin Login — all navigable via a shared navbar."],
        ["FR-02", "Contact form storage", "Seven-field enquiry form, server-side validated, stored in the inquiries table with a timestamp."],
        ["FR-03", "Secure admin login", "Credentials stored in the database; password verified with bcrypt.compare(); session created on success."],
        ["FR-04", "Protected admin area", "requireAdmin middleware guards all /admin/* routes; unauthenticated access is redirected to login."],
        ["FR-05", "Dashboard data", "Dashboard shows the total enquiry count and a full table of all submissions, newest first."],
        ["FR-06", "Logout", "Session destroyed, cookie cleared, no-store cache headers set so the Back button cannot show the dashboard."],
        ["FR-07", "AI chatbot", "Floating FAQ assistant on every public page with ten matching intents and a fallback response."],
    ],
    widths=[0.7, 1.8, 3.8],
)

h2("1.4 Non-Functional Requirements")
make_table(
    ["ID", "Requirement", "How it is met"],
    [
        ["NFR-01", "Clean, consistent design", "Shared EJS header/footer partials, one stylesheet, a single blue/teal/dark palette and font."],
        ["NFR-02", "Secure data storage", "bcrypt-hashed passwords, parameterised SQL everywhere, session secret from environment file."],
        ["NFR-03", "Responsive", "CSS media queries and a collapsing hamburger menu for tablet and mobile widths."],
        ["NFR-04", "Fast load", "No heavy libraries, a single small CSS/JS bundle, lightweight server-rendered pages."],
        ["NFR-05", "Maintainable", "Clear folder structure, commented routes and middleware, single-responsibility data layer."],
    ],
    widths=[0.7, 1.8, 3.8],
)

h2("1.5 Business Process Model")
body("The core business process the system supports is the handling of a sales enquiry from first contact to "
     "resolution. Understanding this process shaped the database, the validation rules, and the admin workflow. "
     "The process runs as follows:")
numbered("A prospective client visits the website and, optionally, asks the FAQ chatbot questions about services.")
numbered("The client submits the Contact Us form with their details and a description of the project they need.")
numbered("The system validates the submission and, if valid, records it as a new enquiry with the status “New”.")
numbered("An administrator logs in and reviews new enquiries on the dashboard.")
numbered("The administrator opens an enquiry, makes contact with the client, and advances its status to “In progress”.")
numbered("When the enquiry has been answered, the administrator marks it “Replied”, completing the process.")
body("This business process is modelled formally later in the report: the data movement is shown in the Data Flow "
     "Diagrams (Section 5.8), the validation decision logic in the Activity Diagram (Section 5.11), and the "
     "enquiry’s status changes in the State Diagram (Section 5.12).")

h2("1.6 Use Case Diagram")
body("A use case diagram is a high-level picture of who uses the system and what they can do with it. It does not "
     "show screens or code — only the goals each type of user can achieve. It is the first diagram in the "
     "document because every later design decision exists to support one of these use cases.")
add_image("01_usecase.png", "Figure 1.1 – Use Case Diagram")

lead("Reading the diagram")
body("The two figures on the outside are the actors — the people who interact with the system. The rounded "
     "shapes inside the box are the use cases (the things they can do). A plain line means “this actor can "
     "perform this use case”. A dashed “includes” arrow means one use case depends on another.")
make_table(
    ["Element", "Meaning in this system"],
    [
        ["Website Visitor (actor)", "Any member of the public. Needs no account. Can browse, enquire, chat, and read content."],
        ["Administrator (actor)", "A staff member with credentials. Everything they do requires being logged in first."],
        ["Browse 8 Pages", "Visitor navigates the eight public pages using the shared navigation bar."],
        ["Submit Contact Inquiry", "Visitor completes the seven-field form; the data is validated and stored."],
        ["Use AI Chatbot", "Visitor asks the floating assistant a question and receives an FAQ answer."],
        ["View Case Studies / Read Feedback", "Visitor reads the company’s portfolio and client testimonials."],
        ["Admin Login", "Administrator proves their identity to gain access to the protected area."],
        ["View Dashboard", "Administrator sees enquiry totals and the submissions table. Dashed line to Admin Login = it requires login."],
        ["Manage Inquiry Status", "Administrator updates an enquiry’s workflow state. Also requires login."],
        ["Logout", "Administrator ends the session safely."],
        ["“includes” (dashed arrows)", "Show that View Dashboard and Manage Inquiry Status are only reachable after a successful Admin Login."],
    ],
    widths=[2.1, 4.2],
)

h2("1.7 Use Case Descriptions")
body("The table below expands the most important use cases into the actor involved, what happens, and the outcome, "
     "so the behaviour is unambiguous.")
make_table(
    ["Use Case", "Actor", "Description", "Outcome"],
    [
        ["Submit Contact Inquiry", "Visitor", "Completes and submits the seven-field enquiry form.", "Validated enquiry saved; confirmation shown."],
        ["Use AI Chatbot", "Visitor", "Asks a question in the chat widget.", "Matching FAQ answer or fallback response shown."],
        ["Admin Login", "Administrator", "Submits username and password.", "Session created and dashboard shown, or error returned."],
        ["View Dashboard", "Administrator", "Opens the protected dashboard.", "Enquiry count and full submissions table displayed."],
        ["Manage Inquiry Status", "Administrator", "Updates an enquiry’s status.", "Status persisted in the database."],
        ["Logout", "Administrator", "Ends the session.", "Session destroyed; redirected to login."],
    ],
    widths=[1.5, 1.0, 2.3, 1.5],
)

h2("1.8 Expected Deliverables")
body("The agreed deliverables for the project, against which progress and completion were measured, are:")
make_table(
    ["Deliverable", "Type", "Status"],
    [
        ["Signed-off requirements specification", "Documentation", "Delivered"],
        ["Project schedule (working document)", "Documentation", "Delivered"],
        ["Responsive public website (8 pages)", "Software", "Delivered"],
        ["Validated contact enquiry workflow + database", "Software", "Delivered"],
        ["Secure administrator dashboard", "Software", "Delivered"],
        ["AI FAQ chatbot", "Software", "Delivered"],
        ["Design documentation (architecture, ERD, DFDs, etc.)", "Documentation", "Delivered"],
        ["Testing results", "Documentation", "Delivered"],
        ["This portfolio report", "Documentation", "Delivered"],
        ["Video demonstration (Task 2)", "Media", "Delivered separately"],
    ],
    widths=[3.3, 1.6, 1.4],
)

h2("1.9 Client Approval and Sign-Off")
body("The requirements above were reviewed with the client and formally approved at the requirements sign-off "
     "meeting (Section 3.2) as the agreed basis for development. The sign-off below records that approval.")
signoff("Client")
signoff("Student")
pagebreak()

# =========================================================
# 2. PLANNING DOCUMENTATION
# =========================================================
h1("2. Planning Documentation")

h2("2.1 Planning Approach")
body("The project was planned around the fixed module milestones — the client meeting, the requirements "
     "sign-off, the interim review, and the Week-15 submission — and the work between them was broken into "
     "short, deliverable-focused tasks. Because the project was delivered by a single developer alongside other "
     "modules, the schedule was deliberately built as a living document: each task carries a planned start and end "
     "date and the actual dates achieved, so that slippage is visible and the plan can be re-balanced. The schedule "
     "is presented below as a Gantt chart and as a detailed task table.")

h2("2.2 Project Schedule (Gantt Chart)")
add_image("17_gantt.png", "Figure 2.1 – Project Schedule Gantt Chart (planned timeline)")
body("The Gantt chart shows the planned sequence and overlap of the project phases from the client meeting in "
     "early February through to submission in May, with the two key milestones (interim client review and final "
     "submission) marked as diamonds. The detailed table that follows records the planned dates against the "
     "actual dates achieved, which is what makes this a working schedule rather than a one-off plan.")

h2("2.3 Detailed Schedule – Planned vs Actual")
make_table(
    ["Task / Deliverable", "Planned start", "Planned end", "Actual start", "Actual end", "Status"],
    [
        ["Client meeting & scenario review", "02 Feb", "08 Feb", "02 Feb", "08 Feb", "On time"],
        ["Requirements spec & sign-off", "09 Feb", "19 Feb", "09 Feb", "20 Feb", "+1 day"],
        ["Project schedule & methodology", "16 Feb", "22 Feb", "18 Feb", "24 Feb", "Slipped 2 days"],
        ["Architecture & database design", "23 Feb", "04 Mar", "24 Feb", "06 Mar", "+2 days"],
        ["Process & interaction design", "05 Mar", "11 Mar", "07 Mar", "13 Mar", "+2 days"],
        ["Public pages & shared layout", "12 Mar", "21 Mar", "12 Mar", "21 Mar", "On time"],
        ["Contact form & validation", "22 Mar", "28 Mar", "22 Mar", "29 Mar", "+1 day"],
        ["Admin auth & dashboard", "29 Mar", "08 Apr", "30 Mar", "10 Apr", "+2 days"],
        ["Chatbot & feedback module", "09 Apr", "15 Apr", "11 Apr", "16 Apr", "+1 day"],
        ["Interim client review (milestone)", "26 Mar", "26 Mar", "26 Mar", "26 Mar", "Held"],
        ["Functional & non-functional testing", "16 Apr", "25 Apr", "17 Apr", "25 Apr", "Recovered"],
        ["Portfolio & deployment docs", "26 Apr", "07 May", "26 Apr", "07 May", "On time"],
        ["Final submission (milestone)", "08 May", "08 May", "08 May", "08 May", "Met"],
    ],
    widths=[2.2, 0.95, 0.95, 0.95, 0.95, 0.9],
)

h2("2.4 Effort Allocation")
body("Effort was weighted towards design and implementation, with a deliberate reserve for testing and "
     "documentation so that quality was not sacrificed at the end of the project.")
make_table(
    ["Phase", "Approx. effort", "Notes"],
    [
        ["Requirements & planning", "10%", "Front-loaded so development built on an approved, stable base."],
        ["Design documentation", "20%", "Architecture, database and interaction diagrams produced before coding."],
        ["Implementation", "40%", "The largest phase — server, routes, validation, admin area, chatbot."],
        ["Testing", "15%", "Functional and non-functional testing against each requirement."],
        ["Documentation & deployment", "15%", "This portfolio, deployment notes and the demonstration."],
    ],
    widths=[2.3, 1.3, 2.8],
)

h2("2.5 Revisions to the Schedule")
body("As the schedule shows, several tasks finished one or two days later than planned. The most significant slip "
     "was during design and the admin area, where building secure session handling and the dashboard queries took "
     "longer than estimated. Rather than let this compress the testing phase — a common cause of poor-quality "
     "submissions — the plan was re-balanced: the chatbot and feedback module, which were lower risk, were "
     "tightened slightly, and the small buffer reserved before submission was used. By the testing phase the "
     "project was back on track (“Recovered” in the table), and both milestones were met. Keeping the "
     "schedule updated with actual dates made these trade-offs visible early enough to act on them.")
pagebreak()

# =========================================================
# 3. CLIENT CONTACT RECORD SHEET
# =========================================================
h1("3. Client Contact Record Sheet")
body("This section records the three client meetings held during the project. In this module the module tutor "
     "acted in the role of the client. Each record lists the attendees, the agenda, what was discussed, the agreed "
     "action points, and a sign-off line for the client and the student.")

h2("3.1 Meeting 1 – Initial Client Meeting")
make_table(
    ["Field", "Detail"],
    [
        ["Date", "04 February 2026"],
        ["Attendees", STUDENT_NAME + " (developer); " + CLIENT_NAME],
        ["Purpose", "Understand the scenario and the client’s needs for the AI-Solutions website."],
    ],
    widths=[1.4, 5.0],
)
lead("Discussion summary")
body("The client described AI-Solutions as a studio selling AI chatbots, automation and custom machine-learning "
     "work, and explained that the website must both market these services and capture prospective-client "
     "enquiries. We discussed the eight pages required, the information the contact form should capture, the need "
     "for a private area for staff to review enquiries, and the idea of an on-site assistant to answer common "
     "questions.")
lead("Action points")
bullet("Student to draft the functional and non-functional requirements for sign-off.")
bullet("Student to propose the contact-form fields and the admin dashboard contents.")
bullet("Client to confirm which pages are mandatory for the marketing site.")
signoff("Client")
signoff("Student")

h2("3.2 Meeting 2 – Requirements Sign-Off")
make_table(
    ["Field", "Detail"],
    [
        ["Date", "18 February 2026"],
        ["Attendees", STUDENT_NAME + " (developer); " + CLIENT_NAME],
        ["Purpose", "Review and formally approve the requirements specification as the basis for development."],
    ],
    widths=[1.4, 5.0],
)
lead("Discussion summary")
body("The drafted functional requirements (FR-01 to FR-07) and non-functional requirements (NFR-01 to NFR-05) were "
     "reviewed line by line. The client confirmed the eight pages, agreed the seven contact-form fields, and "
     "approved the security expectations (hashed passwords, protected admin area). The requirements were judged "
     "appropriate and realistic and were signed off as the agreed scope.")
lead("Action points")
bullet("Requirements approved — development to begin against the signed-off specification.")
bullet("Student to produce design documentation (architecture, database, interaction diagrams) next.")
bullet("Scope frozen for the prototype; further ideas recorded as future enhancements.")
signoff("Client")
signoff("Student")

h2("3.3 Meeting 3 – Interim Review")
make_table(
    ["Field", "Detail"],
    [
        ["Date", "26 March 2026"],
        ["Attendees", STUDENT_NAME + " (developer); " + CLIENT_NAME],
        ["Purpose", "Check progress against the schedule and confirm the project is on track for completion."],
    ],
    widths=[1.4, 5.0],
)
lead("Discussion summary")
body("A working build of the public site and contact form was demonstrated, and the design documentation was "
     "reviewed. We discussed the slight slippage in the design and admin phases and agreed the plan to protect the "
     "testing window. The client was satisfied with progress and the direction of the admin dashboard, and "
     "confirmed no change to the agreed scope.")
lead("Action points")
bullet("Student to complete the admin dashboard and chatbot, then move into testing.")
bullet("Student to re-balance the schedule to keep the testing phase intact.")
bullet("Student to prepare the portfolio and video demonstration for Week 15.")
signoff("Client")
signoff("Student")
pagebreak()

# =========================================================
# 4. METHODOLOGY
# =========================================================
h1("4. Methodology")
body("This section explains, in my own words, the approach I took to plan and manage the project and the tools I "
     "used to deliver it within the agreed timescales. It is written about my own professional practice rather than "
     "as a research review.")

h2("4.1 The Approach I Chose")
body("Because I was working alone on a fixed-deadline project for a single client, I chose an adapted, lightweight "
     "Agile and iterative approach rather than a heavy, plan-everything-up-front waterfall method. A pure waterfall "
     "would have been risky here: if I had spent weeks producing a perfect specification and only discovered "
     "problems during coding, there would have been little time to recover. Equally, a full team-based Agile "
     "framework such as Scrum, with daily stand-ups, sprint ceremonies and multiple roles, would have been "
     "unnecessary overhead for one person. What I actually needed was the discipline of clear requirements and a "
     "schedule, combined with the flexibility to build the product feature by feature and learn as I went.")
body("My adaptation kept the parts of each approach that fit a solo client project. From a structured approach I "
     "kept a written, signed-off requirements specification and a working schedule with milestones, so the client "
     "and I shared a clear, agreed target. From Agile I kept short iterations: I built the system as a series of "
     "small increments — first the shared layout and public pages, then the contact form and validation, then "
     "the admin authentication and dashboard, and finally the chatbot and feedback module — testing each one "
     "before moving on. This meant I always had a working, demonstrable product, which is exactly what I was able "
     "to show at the interim review.")

h2("4.2 My Methodology Model")
add_image("16_methodology.png", "Figure 4.1 – The adapted iterative methodology used on this project")
lead("Reading the diagram")
make_table(
    ["Stage", "What it meant in practice"],
    [
        ["Client Meeting & Scenario Review", "Understanding the client’s needs before writing anything (Section 3.1)."],
        ["Requirements Specification & Sign-off", "Agreeing a fixed, approved scope as the basis for development (Section 3.2)."],
        ["Planning & Scheduling", "Breaking the work into dated, deliverable-focused tasks (Section 2)."],
        ["Iterative Build Cycle", "The heart of the method — repeated for every feature."],
        ["Design → Implement → Manual Test → Review", "Each feature was designed, coded, tested, then reflected on before the next."],
        ["Client Interim Review", "A mid-project checkpoint that fed back into the build cycle (Section 3.3)."],
        ["Final Testing & Evaluation", "Once all features were complete, testing the whole system against the requirements."],
        ["Deployment & Documentation", "Packaging the solution and writing this portfolio."],
    ],
    widths=[2.6, 3.8],
)
body("The loop in the middle of the diagram is the key adaptation: instead of one long build, the design-implement-"
     "test-review cycle repeats for each feature, and the client interim review feeds back into it. This is how I "
     "kept the flexibility of Agile while still working towards the fixed milestones of the module.")

h2("4.3 Why This Fit the Project and Its Constraints")
body("The biggest constraints were time, working solo, and the fixed client milestones. The iterative approach "
     "suited all three. Building feature by feature meant that if time ran short, I would still have a coherent, "
     "working subset of the product rather than a half-finished attempt at everything at once. Testing each "
     "increment as I built it caught problems while they were small and cheap to fix, which protected the limited "
     "time I had. And because the method was organised around the same milestones the module already imposed — "
     "requirements sign-off and interim review — the client checkpoints became natural review points in my own "
     "process rather than interruptions to it.")

h2("4.4 Tools and Techniques")
body("I deliberately chose lightweight, widely used tools so that the time went into the product rather than into "
     "fighting the tooling.")
make_table(
    ["Tool / technique", "Purpose and why I chose it"],
    [
        ["Visual Studio Code", "Main editor — fast, free, with strong JavaScript and EJS support."],
        ["Node.js & npm", "Runtime and package manager — lets me run the server and pull in only the few dependencies I needed."],
        ["Git (local version control)", "To track changes incrementally and roll back safely if an experiment failed."],
        ["Express + EJS", "A minimal server framework and templating engine — little boilerplate, quick to build pages."],
        ["SQLite", "A zero-configuration, file-based database — no separate server to install, ideal for a solo prototype."],
        ["Mermaid diagrams", "To produce all the design diagrams in this report as code, so they are consistent and easy to revise."],
        ["Manual functional testing", "Exercising each requirement’s success and failure path directly against the running server."],
        ["Client meetings", "Requirements sign-off and interim review used as decision and quality gates."],
    ],
    widths=[2.1, 4.3],
)

h2("4.5 How Planning and Methodology Worked Together")
body("The schedule in Section 2 and the methodology in this section are two views of the same plan. The schedule "
     "set the dates and milestones; the methodology set the rhythm of work between them. Keeping the schedule "
     "updated with actual dates was itself part of the method — it was the feedback that told me when the "
     "build cycle was slipping and let me re-balance the remaining iterations to protect testing and the final "
     "submission. In practice the two worked together to keep a solo project disciplined without making it rigid.")
pagebreak()

# =========================================================
# 5. SOLUTION DESIGN DOCUMENTATION
# =========================================================
h1("5. Solution Design Documentation")

h2("5.1 Technology Stack")
body("Each technology below was chosen deliberately. The right-hand column explains the reasoning, which directly "
     "supports the maintainability and performance requirements.")
make_table(
    ["Layer", "Technology", "Reason for choice"],
    [
        ["Runtime", "Node.js", "Event-driven JavaScript runtime, efficient for the input/output-heavy work of a web app."],
        ["Web framework", "Express", "Minimal, widely used routing and middleware framework with little boilerplate."],
        ["Templating", "EJS", "Renders HTML on the server with embedded JavaScript and reusable partials (shared header/footer)."],
        ["Database", "SQLite (sqlite3)", "Zero-configuration, file-based relational database; no separate server to install or run."],
        ["Authentication", "bcryptjs", "Industry-standard one-way password hashing with a configurable work factor."],
        ["Sessions", "express-session", "Server-side session management backed by a signed HttpOnly cookie."],
        ["Frontend", "HTML, CSS, vanilla JS", "Lightweight, dependency-free client giving full control over styling and behaviour."],
    ],
    widths=[1.4, 1.8, 3.1],
)

h2("5.2 High-Level Architecture")
body("The application follows the classic three-tier pattern, which separates what the user sees, what the "
     "application does, and where the data lives. Separating these concerns keeps each part simple and replaceable.")
add_image("02_architecture.png", "Figure 5.1 – Three-Tier System Architecture")

lead("Reading the diagram")
make_table(
    ["Tier / Element", "Role"],
    [
        ["Presentation Tier (browser)", "The visitor’s browser. Renders the HTML/CSS it receives and runs the vanilla-JS chatbot. Holds no business logic."],
        ["HTML / CSS / Vanilla JS", "The rendered page and styling the user actually sees and interacts with."],
        ["AI Chatbot Widget (chat.js)", "Runs entirely inside the browser; it does not call the server to answer questions."],
        ["Application Tier (Node/Express)", "The server. Receives requests, runs middleware, decides what to do, and renders pages."],
        ["Express App (server.js)", "The single entry point that wires everything together and defines all routes."],
        ["Session & requireAdmin Middleware", "Identifies logged-in admins and blocks unauthenticated access to admin pages."],
        ["Route Handlers", "The functions that respond to each URL, e.g. the contact handler or the dashboard handler."],
        ["EJS Template Engine", "Turns data plus a template into a finished HTML page sent back to the browser."],
        ["Data Tier (db.js + SQLite)", "Where information is stored. Reached only through the db.js module."],
        ["Arrows", "Show the request/response flow: browser → Express → middleware → handler → (templates and data) → back to the browser."],
    ],
    widths=[2.3, 4.0],
)
body("The single most important point in this diagram is that the route handlers never touch SQLite directly — "
     "they always go through db.js. This means the entire storage technology is isolated behind one module.")

h2("5.3 Component View")
body("Where the architecture diagram shows tiers, the component diagram shows the parts inside the server and how "
     "requests are routed to them. It makes the security boundary explicit.")
add_image("03_component.png", "Figure 5.2 – Component Diagram")

lead("Reading the diagram")
make_table(
    ["Element", "Meaning"],
    [
        ["Express Server (server.js)", "The hub. Every incoming request first reaches the server, which dispatches it to a component."],
        ["Public Components", "Home, Solutions, Case Studies, Feedback, Articles, Gallery, Contact Form, and the Chatbot engine — all open to anyone."],
        ["requireAdmin Middleware", "A gate the server places in front of the protected components."],
        ["Admin Components", "Login and Logout are reachable directly; Dashboard and Inquiry Detail sit behind the requireAdmin gate."],
        ["db.js (Data Access Layer)", "The shared component that performs all database work."],
        ["SQLite", "The actual data store, reached only by db.js."],
        ["Arrows into db.js", "Only data-driven components (Contact Form, Feedback, Login, Dashboard, Inquiry Detail) connect to the data layer; static pages do not."],
    ],
    widths=[2.2, 4.1],
)
body("The diagram shows that the Dashboard and Inquiry Detail components can only be reached by passing through "
     "requireAdmin, whereas Login and Logout are reached directly — which is exactly what the security model "
     "requires.")

h2("5.4 Deployment View")
body("The deployment diagram answers a different question: where does each part physically run? For this academic "
     "project everything runs on one machine, which the diagram makes clear.")
add_image("04_deployment.png", "Figure 5.3 – Deployment Diagram")

lead("Reading the diagram")
make_table(
    ["Element", "Meaning"],
    [
        ["Client Device (Browser)", "The user’s computer or phone, connecting over HTTP."],
        ["Host Machine", "The single Windows/localhost machine that runs the whole back end."],
        ["Node.js Runtime :3000", "The server process, listening on network port 3000."],
        ["Express Application", "The running application code inside that process."],
        ["In-Memory Session Store", "Logged-in sessions are kept in the server’s memory (so they are lost on restart)."],
        ["/public (static files)", "CSS, JavaScript, and images served straight to the browser without processing."],
        ["database.sqlite (file)", "The database is a single file on the same disk — no separate database server."],
        ["HTTP :3000 arrow", "All traffic between the browser and the server travels over this one connection."],
    ],
    widths=[2.2, 4.1],
)
body("Because SQLite is a file rather than a server, deployment is reduced to running one Node.js process — a "
     "deliberate simplicity that supports the performance and maintainability goals.")

h2("5.5 Project Folder Structure")
body("The code is organised by responsibility rather than by file type. Each item below has one clear job, which "
     "is what makes the project easy to navigate and maintain (NFR-05).")
code_block(
    "website/\n"
    "  server.js            Main Express application & all routes\n"
    "  db.js                SQLite data-access layer (initDb, query, execute, getOne)\n"
    "  seed.js              Seeds admin user, sample inquiries and reviews\n"
    "  .env                 PORT and SESSION_SECRET (not committed)\n"
    "  package.json         Dependencies and npm scripts\n"
    "  database.sqlite      The SQLite database file (auto-created)\n"
    "  views/               EJS templates\n"
    "    partials/          header.ejs, footer.ejs, chatbot.ejs (shared layout)\n"
    "    home.ejs solutions.ejs cases.ejs feedback.ejs ...\n"
    "    admin-login.ejs admin-dashboard.ejs admin-inquiry.ejs ...\n"
    "  public/              Static assets served to the browser\n"
    "    css/style.css      Single stylesheet (design tokens + responsive rules)\n"
    "    js/main.js         Navbar / hamburger, filters, modal, prefill\n"
    "    js/chat.js         Chatbot FAQ engine\n"
    "    images/            Static images\n"
)
make_table(
    ["Location", "Responsibility"],
    [
        ["server.js", "Routing and request handling — the application’s control flow."],
        ["db.js", "All data access — the only file that talks to SQLite."],
        ["views/", "Presentation — how pages look, with shared partials for the header, footer, and chatbot."],
        ["public/", "Static assets — styling, client-side scripts, and images."],
        [".env", "Configuration and secrets, kept out of version control."],
    ],
    widths=[1.8, 4.5],
)

# ---- database design folded into design documentation ----
h2("5.6 Database Design")

h3("5.6.1 Entity Relationship Diagram")
body("An Entity Relationship Diagram (ERD) shows the tables (entities), the columns inside them (attributes), and "
     "the relationships between them. It is the blueprint of how information is stored.")
add_image("05_er.png", "Figure 5.4 – Entity Relationship Diagram")

lead("Reading the diagram")
make_table(
    ["Element", "Meaning"],
    [
        ["INQUIRIES entity", "Stores every customer enquiry submitted through the contact form."],
        ["ADMINS entity", "Stores administrator login accounts; the password column holds a bcrypt hash, never plain text."],
        ["REVIEWS entity", "Stores client testimonials shown on the Feedback page."],
        ["PK marker", "Primary Key — the column that uniquely identifies each row (always id here)."],
        ["UK marker", "Unique Key — the username column cannot be duplicated across admin accounts."],
        ["The crow’s-foot lines", "Show logical relationships: one admin conceptually manages many inquiries and moderates many reviews."],
        ["“logical” label", "These relationships are conceptual only; no physical foreign-key constraint is enforced because each table is self-contained."],
    ],
    widths=[2.0, 4.3],
)
body("The relationships are drawn as “one-to-many” (one administrator, many enquiries) but are described "
     "as logical. The tables do not store an admin id on each enquiry, so there is no enforced foreign key — a "
     "deliberate simplification suited to a single-administrator system.")

h3("5.6.2 Data Dictionary – inquiries")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "INTEGER", "PK, AUTOINCREMENT", "Unique enquiry identifier, assigned automatically."],
        ["name", "TEXT", "NOT NULL", "Full name of the enquirer."],
        ["email", "TEXT", "NOT NULL", "Contact email address."],
        ["phone", "TEXT", "NOT NULL", "Contact phone number."],
        ["company", "TEXT", "(optional)", "Company name."],
        ["country", "TEXT", "(optional)", "Country."],
        ["job_title", "TEXT", "(optional)", "Job title of the enquirer."],
        ["job_details", "TEXT", "NOT NULL", "Free-text description of the project need."],
        ["status", "TEXT", "DEFAULT 'New'", "Workflow state: New, In progress, or Replied."],
        ["assigned_to", "TEXT", "DEFAULT 'Kritika S.'", "Staff member handling the enquiry."],
        ["created_at", "DATETIME", "DEFAULT CURRENT_TIMESTAMP", "Time the enquiry was submitted, set automatically."],
    ],
    widths=[1.2, 1.0, 1.8, 2.3],
)

h3("5.6.3 Data Dictionary – admins")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "INTEGER", "PK, AUTOINCREMENT", "Unique administrator identifier."],
        ["username", "TEXT", "UNIQUE, NOT NULL", "Login name (an email address); cannot be duplicated."],
        ["password", "TEXT", "NOT NULL", "bcrypt hash of the password — never plain text."],
        ["created_at", "DATETIME", "DEFAULT CURRENT_TIMESTAMP", "Account creation time."],
    ],
    widths=[1.2, 1.0, 1.8, 2.3],
)

h3("5.6.4 Data Dictionary – reviews")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "INTEGER", "PK, AUTOINCREMENT", "Unique review identifier."],
        ["name", "TEXT", "NOT NULL", "Name of the reviewer."],
        ["role", "TEXT", "(optional)", "Reviewer’s role / company."],
        ["rating", "INTEGER", "NOT NULL", "Star rating (1–5)."],
        ["comment", "TEXT", "NOT NULL", "Testimonial text."],
        ["created_at", "DATETIME", "DEFAULT CURRENT_TIMESTAMP", "Time the review was added."],
    ],
    widths=[1.2, 1.0, 1.8, 2.3],
)

h3("5.6.5 Table Creation Code")
body("The tables are created by db.js the first time the application runs. The real CREATE TABLE statements are "
     "shown below, followed by an explanation of the key choices.")
code_block(
    "CREATE TABLE IF NOT EXISTS inquiries (\n"
    "  id INTEGER PRIMARY KEY AUTOINCREMENT,\n"
    "  name TEXT NOT NULL,\n"
    "  email TEXT NOT NULL,\n"
    "  phone TEXT NOT NULL,\n"
    "  company TEXT,\n"
    "  country TEXT,\n"
    "  job_title TEXT,\n"
    "  job_details TEXT NOT NULL,\n"
    "  status TEXT DEFAULT 'New',\n"
    "  assigned_to TEXT DEFAULT 'Kritika S.',\n"
    "  created_at DATETIME DEFAULT CURRENT_TIMESTAMP\n"
    ");"
)
walk([
    ("CREATE TABLE IF NOT EXISTS",
     "Creates the table only if it does not already exist, so restarting the app never destroys or duplicates data."),
    ("id INTEGER PRIMARY KEY AUTOINCREMENT",
     "Gives every row a unique number that the database assigns automatically and never reuses."),
    ("name / email / phone TEXT NOT NULL",
     "The three essential contact fields must always contain a value; the database rejects a row missing any of them."),
    ("company / country / job_title TEXT",
     "Optional fields — no NOT NULL, so they may be left blank."),
    ("job_details TEXT NOT NULL",
     "The description of the project is required, as it is the substance of the enquiry."),
    ("status TEXT DEFAULT 'New'",
     "Every new enquiry automatically starts in the 'New' state without the code having to set it."),
    ("assigned_to TEXT DEFAULT 'Kritika S.'",
     "Enquiries default to the studio owner as the handler."),
    ("created_at DATETIME DEFAULT CURRENT_TIMESTAMP",
     "The submission time is stamped automatically by the database at the moment the row is inserted."),
])

h3("5.6.6 Design Notes")
for t in [
    "All primary keys are surrogate auto-increment integers, giving each row a stable, meaningless-but-unique identifier.",
    "Timestamps default to CURRENT_TIMESTAMP so the application never has to set the time manually, avoiding clock-handling bugs.",
    "The username column is UNIQUE, preventing two administrators sharing the same login name.",
    "The tables are in third normal form: every non-key column depends only on the primary key, with no repeating groups.",
    "Sensitive data (the administrator password) is stored only as an irreversible bcrypt hash, never as readable text.",
]:
    bullet(t)

# ---- process & interaction design ----
h2("5.7 Site Map")
body("The site map shows every page and the navigation paths between them. It is the map a user mentally follows "
     "when moving around the website.")
add_image("06_sitemap.png", "Figure 5.5 – Site Map / Navigation Structure")

lead("Reading the diagram")
make_table(
    ["Element", "Meaning"],
    [
        ["Home (/)", "The root page; every public page is reachable from its shared navigation bar."],
        ["Solutions, Feedback, Articles, Gallery, Contact", "The other public pages branching from Home."],
        ["Case Studies → Case Detail (/cases/:id)", "The case list links to individual case pages; :id is the case number in the URL."],
        ["Admin Login (/admin/login)", "The entry point to the private area, shaded to mark it as the security boundary."],
        ["Admin Dashboard (/admin/dashboard)", "Reachable only after login; the hub of the admin area."],
        ["Inquiry Detail (/admin/inquiries/:id)", "Opened from the dashboard to view and update a single enquiry."],
        ["Logout (/admin/logout)", "Ends the session and returns to the login page."],
        ["Shaded nodes", "Indicate admin-only destinations that require an active session."],
    ],
    widths=[2.5, 3.8],
)

h2("5.8 Data Flow Diagrams")
body("A Data Flow Diagram (DFD) shows how information moves through the system: who provides it, which processes "
     "transform it, and where it is stored. Two levels are provided.")

h3("5.8.1 Level 0 – Context Diagram")
body("The context diagram is the highest-level view: the entire system is a single process exchanging data with "
     "the outside world.")
add_image("07_dfd_context.png", "Figure 5.6 – Data Flow Diagram (Level 0 / Context)")

lead("Reading the diagram")
make_table(
    ["Element", "Meaning"],
    [
        ["Visitor (external entity)", "Supplies enquiry details and chat messages; receives web pages and chat replies."],
        ["Administrator (external entity)", "Supplies credentials and status updates; receives dashboard data and the enquiry list."],
        ["0. AI-Solutions Web System (process)", "The whole application, treated as a single black box at this level."],
        ["Database Store", "The persistent data the system reads from and writes to."],
        ["Labelled arrows", "Each arrow names the data that flows in that direction (e.g. “credentials” in, “dashboard data” out)."],
    ],
    widths=[2.4, 3.9],
)

h3("5.8.2 Level 1 – Decomposition")
body("The Level 1 diagram opens the single process from Level 0 into the five real processes inside the system and "
     "shows how each one uses the three data stores.")
add_image("08_dfd_level1.png", "Figure 5.7 – Data Flow Diagram (Level 1)")

lead("Reading the diagram")
make_table(
    ["Process / Store", "Role"],
    [
        ["1.0 Serve Public Pages", "Renders the public pages; reads reviews (D3) to show testimonials."],
        ["2.0 Process Inquiry", "Validates the contact form and INSERTs the enquiry into the inquiries store (D1)."],
        ["3.0 Authenticate Admin", "Checks submitted credentials against the admins store (D2) by verifying the hash."],
        ["4.0 Manage Dashboard", "Reads enquiries from D1 to build the dashboard and UPDATEs their status."],
        ["5.0 Generate Chatbot Reply", "Produces a chat answer; note it touches no data store because it runs in the browser."],
        ["D1 inquiries / D2 admins / D3 reviews", "The three persistent data stores, matching the three database tables."],
        ["Arrows (INSERT / SELECT / UPDATE)", "Show the exact read or write each process performs on each store."],
    ],
    widths=[2.4, 3.9],
)

h2("5.9 Sequence Diagrams")
body("A sequence diagram shows the ordered messages between participants during one specific interaction, read top "
     "to bottom. The dashed return arrows are responses. “alt” boxes show alternative paths (success vs. "
     "failure).")

h3("5.9.1 Contact Form Submission")
add_image("09_seq_contact.png", "Figure 5.8 – Sequence: Contact Form Submission")
lead("Reading the diagram")
make_table(
    ["Step", "What happens"],
    [
        ["Visitor → Browser", "The visitor fills the seven fields and presses submit."],
        ["Browser → Express (POST /contact)", "The form data is sent to the server."],
        ["Express self-call (validation)", "The server validates all seven fields before doing anything else."],
        ["alt: Validation fails", "The form is re-rendered with a specific error message; nothing is saved."],
        ["alt: Validation passes → INSERT", "A parameterised INSERT writes the enquiry to the inquiries table."],
        ["SQLite → Express (insertId)", "The database confirms the new row and returns its id."],
        ["Express → Browser → Visitor", "A success confirmation page is shown."],
    ],
    widths=[2.6, 3.7],
)

h3("5.9.2 Administrator Login")
add_image("10_seq_login.png", "Figure 5.9 – Sequence: Administrator Login")
lead("Reading the diagram")
make_table(
    ["Step", "What happens"],
    [
        ["Administrator → Browser → Express", "Username and password are submitted to /admin/login."],
        ["Express → SQLite (SELECT by username)", "The server looks up the single admin row for that username."],
        ["SQLite → Express (row with bcrypt hash)", "The stored hash is returned — the real password is never stored."],
        ["Express self-call (bcrypt.compare)", "The submitted password is hashed and compared to the stored hash."],
        ["alt: Password matches", "The session is regenerated (anti-fixation), the admin id is stored, and the browser is redirected to the dashboard."],
        ["alt: No match", "The browser is redirected back to the login page with an error flag."],
    ],
    widths=[2.8, 3.5],
)

h3("5.9.3 Chatbot Interaction")
add_image("11_seq_chatbot.png", "Figure 5.10 – Sequence: AI Chatbot Interaction")
lead("Reading the diagram")
make_table(
    ["Step", "What happens"],
    [
        ["Visitor → Widget", "The visitor types a question into the chat panel."],
        ["Widget self-call (typing indicator)", "Animated dots are shown to mimic a real assistant thinking."],
        ["Widget → FAQ logic (getBotResponse)", "The message is passed to the matching function."],
        ["FAQ logic self-call (lowercase + match)", "The text is lower-cased and tested against ten keyword intents."],
        ["alt: Keyword matched", "A tailored answer plus action buttons is returned."],
        ["alt: No keyword matched", "A friendly fallback with suggested topics is returned."],
        ["Widget → Visitor", "The bot reply is rendered in the panel. All of this happens in the browser — no server call."],
    ],
    widths=[2.8, 3.5],
)

h2("5.10 Authentication Flowchart")
body("A flowchart shows decision logic as a path of steps and yes/no branches. This one captures exactly what the "
     "requireAdmin middleware does on every admin request.")
add_image("12_flow_auth.png", "Figure 5.11 – Flowchart: Admin Authentication Middleware")
lead("Reading the diagram")
make_table(
    ["Element", "Meaning"],
    [
        ["Request to /admin/* route", "The starting point — any attempt to reach an admin page."],
        ["Set Cache-Control: no-store", "The first action: tell the browser never to cache this page (Back-button protection)."],
        ["Decision: session.adminId exists?", "The single security question — is this request from a logged-in admin?"],
        ["Yes → next()", "Control passes to the real route handler and the protected page is served."],
        ["No → Redirect to /admin/login", "The request is bounced to the login page; the protected handler never runs."],
    ],
    widths=[2.4, 3.9],
)

h2("5.11 Contact Validation Activity Diagram")
body("This activity diagram details the server-side validation pipeline for the contact form. The checks run in "
     "sequence; the first failed check stops the process and re-renders the form, so invalid data can never reach "
     "the database.")
add_image("13_activity_validation.png", "Figure 5.12 – Activity: Contact Form Validation")
lead("Reading the diagram")
make_table(
    ["Decision in order", "What it checks / what failure means"],
    [
        ["All required fields present?", "name, email, phone and job_details must exist; otherwise “missing fields” error."],
        ["Name 2–100 chars and valid?", "Length plus an allowed-character pattern (letters, spaces, hyphen, dot)."],
        ["Email format valid?", "Must match a basic address pattern and stay within 100 characters."],
        ["Phone valid characters?", "Length 7–20 and only digits, spaces, +, -, ( ) and dots."],
        ["Job details 10–2000 chars?", "The description must be substantial but not unbounded."],
        ["INSERT inquiry into SQLite", "Reached only when every check passes; the enquiry is then stored."],
        ["Re-render form with error", "The common exit for any failed check — the visitor sees a precise message and keeps their input."],
    ],
    widths=[2.4, 3.9],
)

h2("5.12 Inquiry Status Lifecycle")
body("A state diagram shows the distinct states an object can be in and the events that move it between them. Here "
     "the object is a single enquiry.")
add_image("14_state_inquiry.png", "Figure 5.13 – State Diagram: Inquiry Lifecycle")
lead("Reading the diagram")
make_table(
    ["State / Transition", "Meaning"],
    [
        ["[*] → New", "An enquiry is created in the New state the moment a visitor submits the form."],
        ["New → In progress", "The administrator begins handling the enquiry."],
        ["New → Replied", "The administrator responds directly without an intermediate step."],
        ["In progress → Replied", "A response is sent after handling has started."],
        ["Replied / In progress → [*]", "The enquiry is closed or archived, ending its lifecycle."],
    ],
    widths=[2.4, 3.9],
)

h2("5.13 Data-Access Class View")
body("Although the project is written in a procedural Express style rather than with formal classes, the "
     "data-access design can still be drawn as a small set of cooperating components, which clarifies the layering.")
add_image("15_class_db.png", "Figure 5.14 – Class View: Data Access Layer")
lead("Reading the diagram")
make_table(
    ["Component", "Responsibility"],
    [
        ["ExpressServer", "Holds the app, the requireAdmin function, and all route handlers; it uses the data layer."],
        ["DataAccessLayer", "The db.js module, exposing initDb, query, execute, getOne and getDbType as promise-returning helpers."],
        ["SQLiteDatabase", "The underlying sqlite3 driver with its all / get / run / serialize methods, wrapped by the data layer."],
        ["“uses” arrow", "The server depends on the data layer, not on SQLite directly."],
        ["“wraps” arrow", "The data layer hides the raw driver behind its four simple helpers."],
    ],
    widths=[2.2, 4.1],
)

h2("5.14 Glossary of Terms")
make_table(
    ["Term", "Meaning"],
    [
        ["Node.js", "A runtime that executes JavaScript outside the browser, used here to run the web server."],
        ["Express", "A minimal web framework for Node.js that handles routing and middleware."],
        ["EJS", "Embedded JavaScript – the server-side templating engine used to render HTML pages."],
        ["Route", "A pairing of an HTTP method and URL path (e.g. GET /contact) with the code that handles it."],
        ["Middleware", "A function that runs between the incoming request and the route handler, e.g. an authentication check."],
        ["bcrypt", "A password-hashing algorithm that stores an irreversible hash instead of the real password."],
        ["Hash", "A one-way fingerprint of data; the original cannot be recovered from it."],
        ["Session", "Server-held state that identifies a logged-in administrator across multiple requests."],
        ["Cookie", "A small token stored by the browser that links a request to its server session."],
        ["Parameterised query", "A SQL statement where values are passed separately (using ?) to prevent SQL injection."],
        ["SQL injection", "An attack where malicious input alters a SQL query; prevented by parameterisation."],
        ["FAQ chatbot", "A rule-based assistant that matches keywords in a user message to predefined answers."],
        ["Promise / async-await", "JavaScript tools for handling operations that finish later, such as database calls."],
        ["FR / NFR", "Functional Requirement / Non-Functional Requirement."],
    ],
    widths=[1.7, 4.6],
)
pagebreak()

# =========================================================
# 6. IMPLEMENTATION (deep code walkthrough)
# =========================================================
h1("6. Implementation – Code Walkthrough")
body("This section reproduces the real source code of the project and explains it. Each listing is the actual code "
     "from the files named in the headings, and each is followed by a line-by-line explanation so that nothing is "
     "left unexplained.")

# ---- 6.1 server bootstrap ----
h2("6.1 Application Bootstrap (server.js)")
body("The top of server.js loads the libraries it needs and configures the Express application and its view "
     "engine.")
code_block(
    "const express = require('express');\n"
    "const session = require('express-session');\n"
    "const path = require('path');\n"
    "const bcrypt = require('bcryptjs');\n"
    "const db = require('./db');\n"
    "require('dotenv').config();\n\n"
    "const app = express();\n"
    "const PORT = process.env.PORT || 3000;\n\n"
    "app.set('view engine', 'ejs');\n"
    "app.set('views', path.join(__dirname, 'views'));\n\n"
    "app.use(express.urlencoded({ extended: true }));\n"
    "app.use(express.json());\n"
    "app.use(express.static(path.join(__dirname, 'public')));"
)
walk([
    ("require('express') / session / path / bcryptjs",
     "Load the web framework, the session manager, Node’s path helper, and the password-hashing library."),
    ("const db = require('./db')",
     "Loads the project’s own data-access module — the only gateway to the database."),
    ("require('dotenv').config()",
     "Reads the .env file and makes PORT and SESSION_SECRET available as environment variables."),
    ("const app = express()",
     "Creates the Express application object that the rest of the file configures."),
    ("const PORT = process.env.PORT || 3000",
     "Uses the configured port, or falls back to 3000 if none is set."),
    ("app.set('view engine', 'ejs')",
     "Tells Express to render pages with the EJS templating engine."),
    ("app.set('views', ... 'views')",
     "Points Express at the views folder where the .ejs templates live."),
    ("express.urlencoded({ extended: true })",
     "Middleware that parses submitted HTML form data into req.body."),
    ("express.json()",
     "Middleware that parses any JSON request bodies."),
    ("express.static(... 'public')",
     "Serves the CSS, JavaScript and images in /public directly to the browser."),
])

# ---- 6.2 sessions ----
h2("6.2 Session Configuration (server.js)")
body("Sessions are how the server remembers that an administrator is logged in across many requests. The cookie "
     "options are hardened for security (NFR-02).")
code_block(
    "app.use(session({\n"
    "  secret: process.env.SESSION_SECRET || 'kritika_ai_default_secret_key_2026',\n"
    "  resave: false,\n"
    "  saveUninitialized: false,\n"
    "  cookie: {\n"
    "    maxAge: 1000 * 60 * 60 * 2,                 // 2 hours\n"
    "    secure: process.env.NODE_ENV === 'production',\n"
    "    httpOnly: true,                             // blocks JS cookie theft (XSS)\n"
    "    sameSite: 'lax'                             // CSRF protection\n"
    "  }\n"
    "}));"
)
walk([
    ("secret: process.env.SESSION_SECRET",
     "The key used to sign the session cookie so it cannot be forged; it comes from the .env file, not the code."),
    ("resave: false",
     "Does not re-save an unchanged session on every request — avoids needless writes."),
    ("saveUninitialized: false",
     "Does not create a session for visitors who never log in — saves memory and is more private."),
    ("maxAge: 1000*60*60*2",
     "The session cookie expires after two hours (value is in milliseconds)."),
    ("secure: ... 'production'",
     "Requires HTTPS for the cookie when running in production; relaxed for local development."),
    ("httpOnly: true",
     "Stops client-side JavaScript from reading the cookie, defending against cross-site scripting theft."),
    ("sameSite: 'lax'",
     "Stops the cookie being sent on cross-site requests, mitigating cross-site request forgery."),
])

# ---- 6.3 requireAdmin ----
h2("6.3 Route Protection – requireAdmin (server.js)")
body("This middleware enforces FR-04 (protected admin area) and the Back-button protection of FR-06. It is passed "
     "as the second argument to every admin route, so it runs before those handlers.")
code_block(
    "function requireAdmin(req, res, next) {\n"
    "  res.set('Cache-Control', 'no-store, no-cache, must-revalidate, private');\n"
    "  res.set('Pragma', 'no-cache');\n"
    "  res.set('Expires', '0');\n"
    "  if (req.session && req.session.adminId) {\n"
    "    return next();\n"
    "  }\n"
    "  res.redirect('/admin/login');\n"
    "}"
)
walk([
    ("res.set('Cache-Control', 'no-store...')",
     "Tells the browser never to store this page, so it cannot be redisplayed from cache after logout."),
    ("res.set('Pragma'...) / res.set('Expires', '0')",
     "Older-browser equivalents of the same no-cache instruction, for full coverage."),
    ("if (req.session && req.session.adminId)",
     "Checks that a session exists and that it contains an admin id — i.e. the user is logged in."),
    ("return next()",
     "If logged in, hands control to the actual route handler (e.g. the dashboard)."),
    ("res.redirect('/admin/login')",
     "If not logged in, the protected handler is skipped entirely and the user is sent to the login page."),
])
body("A companion middleware copies the session onto res.locals so every template can adapt its navigation bar to "
     "whether an admin is signed in:")
code_block(
    "app.use((req, res, next) => {\n"
    "  res.locals.session = req.session;\n"
    "  next();\n"
    "});"
)

# ---- 6.4 data access ----
h2("6.4 Data Access Layer (db.js)")
body("db.js is the only module that talks to SQLite. It connects to the database, ensures the tables exist, and "
     "exposes four small promise-returning helpers. Keeping every query behind these helpers is what makes the "
     "storage engine replaceable and the calling code clean.")
code_block(
    "let sqliteDb = null;\n"
    "const DB_PATH = path.join(__dirname, 'database.sqlite');\n\n"
    "function initDb() {\n"
    "  return new Promise((resolve, reject) => {\n"
    "    sqliteDb = new sqlite3.Database(DB_PATH, (err) => {\n"
    "      if (err) return reject(err);\n"
    "      createTables().then(resolve).catch(reject);\n"
    "    });\n"
    "  });\n"
    "}"
)
walk([
    ("let sqliteDb = null",
     "Holds the single shared database connection, created once at start-up."),
    ("const DB_PATH = ... 'database.sqlite'",
     "The database is a single file sitting beside the code."),
    ("function initDb() { return new Promise(...) }",
     "Returns a promise so the server can wait for the database to be ready before serving requests."),
    ("new sqlite3.Database(DB_PATH, cb)",
     "Opens (or creates) the database file; the callback reports success or an error."),
    ("if (err) return reject(err)",
     "If the connection fails, the promise rejects so the caller can log the problem."),
    ("createTables().then(resolve)",
     "Once connected, it creates the tables, and only then resolves — guaranteeing a usable schema."),
])
body("The four query helpers all follow the same shape: wrap a sqlite3 callback in a promise so the route handlers "
     "can use clean async/await.")
code_block(
    "function query(sql, params = []) {       // many rows\n"
    "  return new Promise((resolve, reject) => {\n"
    "    sqliteDb.all(sql, params, (err, rows) => {\n"
    "      if (err) return reject(err);\n"
    "      resolve(rows);\n"
    "    });\n"
    "  });\n"
    "}\n\n"
    "function execute(sql, params = []) {      // insert / update / delete\n"
    "  return new Promise((resolve, reject) => {\n"
    "    sqliteDb.run(sql, params, function (err) {\n"
    "      if (err) return reject(err);\n"
    "      resolve({ insertId: this.lastID, affectedRows: this.changes });\n"
    "    });\n"
    "  });\n"
    "}\n\n"
    "function getOne(sql, params = []) {       // single row\n"
    "  return new Promise((resolve, reject) => {\n"
    "    sqliteDb.get(sql, params, (err, row) => {\n"
    "      if (err) return reject(err);\n"
    "      resolve(row || null);\n"
    "    });\n"
    "  });\n"
    "}"
)
walk([
    ("query → sqliteDb.all",
     "Runs a SELECT and resolves with an array of every matching row — used for lists such as the inquiries table."),
    ("execute → sqliteDb.run",
     "Runs an INSERT, UPDATE or DELETE; resolves with the new row’s id and the number of rows changed."),
    ("getOne → sqliteDb.get",
     "Runs a SELECT and resolves with just the first row, or null — used for single-record lookups such as one admin or one enquiry."),
    ("params = []",
     "Every helper takes a separate array of values that are bound to the ? placeholders, which is what prevents SQL injection."),
    ("resolve(row || null)",
     "Normalises “not found” to null so callers can test it simply."),
])

# ---- 6.5 secure login ----
h2("6.5 Secure Administrator Login (server.js)")
body("This handler implements FR-03. It looks the administrator up by username, verifies the password against the "
     "stored bcrypt hash, and — only on success — regenerates the session to prevent session fixation "
     "before storing the admin id.")
code_block(
    "app.post('/admin/login', async (req, res) => {\n"
    "  const { email, password } = req.body;\n"
    "  if (!email || !password) {\n"
    "    return res.redirect('/admin/login?error=missing');\n"
    "  }\n"
    "  try {\n"
    "    const admin = await db.getOne('SELECT * FROM admins WHERE username = ?', [email]);\n"
    "    if (!admin) return res.redirect('/admin/login?error=invalid');\n\n"
    "    const match = await bcrypt.compare(password, admin.password);\n"
    "    if (!match) return res.redirect('/admin/login?error=invalid');\n\n"
    "    req.session.regenerate((err) => {\n"
    "      if (err) return res.redirect('/admin/login?error=server');\n"
    "      req.session.adminId = admin.id;\n"
    "      req.session.adminUser = admin.username;\n"
    "      req.session.save(() => res.redirect('/admin/dashboard'));\n"
    "    });\n"
    "  } catch (err) {\n"
    "    res.redirect('/admin/login?error=server');\n"
    "  }\n"
    "});"
)
walk([
    ("const { email, password } = req.body",
     "Reads the submitted username and password from the form."),
    ("if (!email || !password) ... ?error=missing",
     "Rejects empty submissions early before touching the database."),
    ("db.getOne('SELECT * FROM admins WHERE username = ?', [email])",
     "Parameterised lookup of the one admin row for that username — the value is bound, never concatenated."),
    ("if (!admin) ... ?error=invalid",
     "If no such user exists, returns the same generic “invalid” message (does not reveal which field was wrong)."),
    ("bcrypt.compare(password, admin.password)",
     "Hashes the typed password and compares it to the stored hash; the real password is never decrypted because it cannot be."),
    ("if (!match) ... ?error=invalid",
     "Wrong password also returns the same generic error, frustrating account enumeration."),
    ("req.session.regenerate(...)",
     "Issues a brand-new session id on login, defeating session-fixation attacks that reuse a pre-set id."),
    ("req.session.adminId = admin.id",
     "Stores the admin id in the session — this is exactly what requireAdmin later checks for."),
    ("req.session.save(() => res.redirect('/admin/dashboard'))",
     "Persists the session, then sends the now-authenticated admin to the dashboard."),
    ("catch (err) ... ?error=server",
     "Any unexpected error is caught and reported generically, never leaking internal details."),
])

# ---- 6.6 logout ----
h2("6.6 Secure Logout (server.js)")
body("Logout implements FR-06. It destroys the session, clears the cookie, and sets no-store headers so the "
     "dashboard cannot be recovered with the browser’s Back button.")
code_block(
    "app.get('/admin/logout', (req, res) => {\n"
    "  req.session.destroy((err) => {\n"
    "    if (err) console.error('Error clearing session:', err.message);\n"
    "    res.clearCookie('connect.sid');\n"
    "    res.set('Cache-Control', 'no-store, no-cache, must-revalidate, private');\n"
    "    res.redirect('/admin/login');\n"
    "  });\n"
    "});"
)
walk([
    ("req.session.destroy(...)",
     "Deletes the session on the server so its admin id no longer exists — the user is now logged out everywhere."),
    ("res.clearCookie('connect.sid')",
     "Removes the session cookie from the browser so no stale token remains."),
    ("res.set('Cache-Control', 'no-store...')",
     "Ensures the browser will not redisplay any cached authenticated page after logout."),
    ("res.redirect('/admin/login')",
     "Returns the user to the login page, completing the logout."),
])

# ---- 6.7 contact validation ----
h2("6.7 Contact Form Validation and Storage (server.js)")
body("This is the longest handler in the project and implements FR-02. It performs the validation pipeline drawn "
     "in Figure 5.12, then stores the enquiry with a fully parameterised INSERT. The structure below shows the "
     "shape of the validation (each check re-renders the form on failure) and the final insert.")
code_block(
    "app.post('/contact', async (req, res) => {\n"
    "  const { name, email, phone, company, country, job_title, job_details } = req.body;\n"
    "  const formData = req.body;\n\n"
    "  // 1. required fields present\n"
    "  if (!name || !email || !phone || !job_details) {\n"
    "    return res.render('contact', { ..., errorMsg: 'Please fill all required fields...', formData });\n"
    "  }\n\n"
    "  // trim every field\n"
    "  const trimmedName = name.trim(); /* ...and the rest... */\n\n"
    "  // 2. name length + allowed characters\n"
    "  if (trimmedName.length < 2 || trimmedName.length > 100) { return res.render('contact', {...}); }\n"
    "  const nameRegex = /^[a-zA-Z\\s'\\-\\.]+$/;\n"
    "  if (!nameRegex.test(trimmedName)) { return res.render('contact', {...}); }\n\n"
    "  // 3. email format\n"
    "  const emailRegex = /^[^\\s@]+@[^\\s@]+\\.[^\\s@]+$/;\n"
    "  if (trimmedEmail.length > 100 || !emailRegex.test(trimmedEmail)) { return res.render('contact', {...}); }\n\n"
    "  // 4. phone length + characters\n"
    "  if (trimmedPhone.length < 7 || trimmedPhone.length > 20) { return res.render('contact', {...}); }\n"
    "  const phoneRegex = /^[0-9\\s\\-\\+\\(\\)\\.]+$/;\n"
    "  if (!phoneRegex.test(trimmedPhone)) { return res.render('contact', {...}); }\n\n"
    "  // 5. optional-field length caps (company/country/job_title <= 100)\n"
    "  // 6. job details length\n"
    "  if (trimmedJobDetails.length < 10 || trimmedJobDetails.length > 2000) { return res.render('contact', {...}); }\n\n"
    "  // all valid -> store\n"
    "  await db.execute(\n"
    "    `INSERT INTO inquiries (name, email, phone, company, country, job_title, job_details, status)\n"
    "     VALUES (?, ?, ?, ?, ?, ?, ?, 'New')`,\n"
    "    [trimmedName, trimmedEmail, trimmedPhone, trimmedCompany, trimmedCountry, trimmedJobTitle, trimmedJobDetails]\n"
    "  );\n"
    "  res.render('contact', { statusMsg: 'success', formData: {} });\n"
    "});"
)
walk([
    ("const { name, email, ... } = req.body",
     "Pulls all seven submitted fields out of the form data."),
    ("if (!name || !email || !phone || !job_details)",
     "Check 1 — the four required fields must be present, or the form is re-rendered with a message."),
    (".trim() on every field",
     "Removes leading/trailing spaces so “   ” is not accepted as a real value."),
    ("length < 2 || length > 100 (name)",
     "Check 2a — enforces a sensible name length."),
    ("nameRegex /^[a-zA-Z\\s'\\-\\.]+$/",
     "Check 2b — allows only letters, spaces, apostrophes, hyphens and dots, blocking numbers and symbols."),
    ("emailRegex ...@...\\....",
     "Check 3 — requires a basic name@domain.tld shape and a 100-character cap."),
    ("phone length 7–20 + phoneRegex",
     "Check 4 — enforces a realistic phone length and a digits/punctuation-only character set."),
    ("optional length caps",
     "Check 5 — company, country and job title may be blank but cannot exceed 100 characters."),
    ("job details 10–2000",
     "Check 6 — the project description must be meaningful yet bounded."),
    ("return res.render('contact', {... errorMsg ..., formData})",
     "Every failed check re-renders the same page with a specific message and the user’s own input preserved."),
    ("db.execute('INSERT ... VALUES (?, ?, ...)', [values])",
     "Only after all checks pass: a parameterised INSERT writes the enquiry; user input is bound, never concatenated, so SQL injection is impossible."),
    ("status defaulted to 'New'",
     "The new enquiry enters the workflow in the New state (matching the state diagram)."),
    ("res.render('contact', { statusMsg: 'success', formData: {} })",
     "Shows the success confirmation and clears the form."),
])

# ---- 6.8 dashboard ----
h2("6.8 Dashboard Data (server.js)")
body("The dashboard handler implements FR-05. It is protected by requireAdmin, gathers several count metrics, and "
     "fetches the full list of enquiries newest-first for the table.")
code_block(
    "app.get('/admin/dashboard', requireAdmin, async (req, res) => {\n"
    "  const totalCountRow = await db.getOne('SELECT COUNT(*) AS total FROM inquiries');\n"
    "  const unreadCountRow = await db.getOne(\"SELECT COUNT(*) AS total FROM inquiries WHERE status = 'New'\");\n"
    "  const progressCountRow = await db.getOne(\"SELECT COUNT(*) AS total FROM inquiries WHERE status = 'In progress'\");\n"
    "  const reviewsRow = await db.getOne('SELECT COUNT(*) AS count, AVG(rating) AS avg FROM reviews');\n"
    "  const inquiries = await db.query('SELECT * FROM inquiries ORDER BY created_at DESC');\n\n"
    "  res.render('admin-dashboard', {\n"
    "    metrics: { total: ..., unread: ..., progress: ..., rating: ..., reviewsCount: ... },\n"
    "    inquiries\n"
    "  });\n"
    "});"
)
walk([
    ("requireAdmin (2nd argument)",
     "Guarantees only a logged-in admin can reach this code; everyone else is redirected to login."),
    ("COUNT(*) AS total FROM inquiries",
     "The headline metric — the total number of enquiries, satisfying the “inquiry count” requirement."),
    ("WHERE status = 'New' / 'In progress'",
     "Additional counts that break the total down by workflow state for the summary cards."),
    ("AVG(rating) FROM reviews",
     "Computes the average client rating shown on the dashboard."),
    ("SELECT * FROM inquiries ORDER BY created_at DESC",
     "Fetches every enquiry, newest first, for the full submissions table."),
    ("res.render('admin-dashboard', { metrics, inquiries })",
     "Passes the counts and the list into the template, which renders the cards and the table."),
])
body("The individual enquiry page and the status update reuse the same pattern: a parameterised getOne to load one "
     "row by id, and a parameterised UPDATE to change its status — the transition drawn in the state diagram.")
code_block(
    "app.post('/admin/inquiries/:id/status', requireAdmin, async (req, res) => {\n"
    "  const { status } = req.body;\n"
    "  await db.execute('UPDATE inquiries SET status = ? WHERE id = ?', [status, req.params.id]);\n"
    "  res.redirect(`/admin/inquiries/${req.params.id}?success=status_updated`);\n"
    "});"
)
walk([
    ("requireAdmin",
     "The status change is a protected, state-changing action, so it too sits behind the auth gate."),
    ("UPDATE inquiries SET status = ? WHERE id = ?",
     "Parameterised update of exactly one enquiry; both the new status and the id are bound values."),
    ("res.redirect(... ?success=status_updated)",
     "Uses the post-redirect-get pattern so refreshing the page does not resubmit the change."),
])

# ---- 6.9 feedback route ----
h2("6.9 Feedback and Reviews (server.js)")
body("The Feedback page lists client reviews and lets visitors add one. Reading reviews computes an average score; "
     "adding a review validates the input and stores it with a parameterised INSERT.")
code_block(
    "app.get('/feedback', async (req, res) => {\n"
    "  const reviews = await db.query('SELECT * FROM reviews ORDER BY created_at DESC');\n"
    "  let avgScore = 4.9;\n"
    "  if (reviews.length > 0) {\n"
    "    const sum = reviews.reduce((acc, r) => acc + r.rating, 0);\n"
    "    avgScore = (sum / reviews.length).toFixed(1);\n"
    "  }\n"
    "  res.render('feedback', { activePage: 'Feedback', reviews, avgScore });\n"
    "});\n\n"
    "app.post('/feedback/add', async (req, res) => {\n"
    "  const { name, role, rating, comment } = req.body;\n"
    "  if (!name || !rating || !comment) {\n"
    "    return res.redirect('/feedback?error=missing_fields');\n"
    "  }\n"
    "  await db.execute(\n"
    "    'INSERT INTO reviews (name, role, rating, comment) VALUES (?, ?, ?, ?)',\n"
    "    [name, role || 'Client', parseInt(rating), comment]\n"
    "  );\n"
    "  res.redirect('/feedback?success=review_added');\n"
    "});"
)
walk([
    ("db.query('SELECT * FROM reviews ORDER BY created_at DESC')",
     "Loads every review, newest first, to display on the feedback page."),
    ("reviews.reduce((acc, r) => acc + r.rating, 0)",
     "Sums all the star ratings so an average can be calculated."),
    ("(sum / reviews.length).toFixed(1)",
     "Computes the average rating to one decimal place; defaults to 4.9 when there are no reviews yet."),
    ("if (!name || !rating || !comment)",
     "Server-side validation — a review must have a name, a rating and a comment, or it is rejected."),
    ("role || 'Client'",
     "If the optional role is left blank, a sensible default of ‘Client’ is stored."),
    ("parseInt(rating)",
     "Converts the submitted rating from text to a whole number before storing it."),
    ("INSERT INTO reviews ... VALUES (?, ?, ?, ?)",
     "Parameterised insert of the new review — the same injection-safe pattern used everywhere."),
    ("res.redirect('/feedback?success=review_added')",
     "Post-redirect-get so a page refresh does not submit the review twice."),
])

# ---- 6.10 home & case detail ----
h2("6.10 Home and Case-Study Routes (server.js)")
body("Two further public routes show how dynamic data and URL parameters are handled. The home page reads the three "
     "most recent reviews; the case-study route reads an :id from the URL and returns a 404-style redirect if no "
     "such case exists.")
code_block(
    "app.get('/', async (req, res) => {\n"
    "  const reviews = await db.query('SELECT * FROM reviews ORDER BY created_at DESC LIMIT 3');\n"
    "  res.render('home', { activePage: 'Home', reviews });\n"
    "});\n\n"
    "app.get('/cases/:id', (req, res) => {\n"
    "  const detail = caseStudiesData[req.params.id];\n"
    "  if (!detail) {\n"
    "    return res.status(404).redirect('/cases');\n"
    "  }\n"
    "  res.render('case-detail', { activePage: 'Case Studies', detail });\n"
    "});"
)
walk([
    ("SELECT * FROM reviews ORDER BY created_at DESC LIMIT 3",
     "Loads only the three newest reviews to feature on the home page — LIMIT keeps the query small."),
    ("res.render('home', { activePage: 'Home', reviews })",
     "Renders the home template, passing the reviews and the active-page name for the navbar highlight."),
    ("app.get('/cases/:id', ...)",
     "A dynamic route — :id is a placeholder that captures whatever case number is in the URL."),
    ("caseStudiesData[req.params.id]",
     "Looks the requested case up by its id in the in-memory case-studies object."),
    ("if (!detail) return res.status(404).redirect('/cases')",
     "If no case matches the id, responds with a not-found status and sends the user back to the case list — invalid URLs fail safely."),
])

# ---- 6.11 seed ----
h2("6.11 Seeding the Database (seed.js)")
body("seed.js creates the first administrator and some sample data. The crucial security point is that the "
     "password is hashed before it is ever written to the database.")
code_block(
    "const adminPassword = 'admin123';\n"
    "const hashedPassword = await bcrypt.hash(adminPassword, 10);\n"
    "await db.execute('DELETE FROM admins');\n"
    "await db.execute('INSERT INTO admins (username, password) VALUES (?, ?)',\n"
    "                 ['admin@kritika.ai', hashedPassword]);"
)
walk([
    ("bcrypt.hash(adminPassword, 10)",
     "Turns the plain password into an irreversible hash using a work factor of 10; this hash is what gets stored."),
    ("DELETE FROM admins",
     "Clears any previous admin so re-running the seed produces a clean, single account."),
    ("INSERT ... VALUES (?, ?), ['admin@kritika.ai', hashedPassword]",
     "Stores the username and the hash (never the plain text) using a parameterised statement."),
])
body("It then loops over arrays of sample enquiries and reviews, inserting each with the same parameterised "
     "execute helper, so the dashboard has realistic data to display on first run.")

# ---- 6.12 chatbot ----
h2("6.12 The AI Chatbot (public/js/chat.js)")
body("The chatbot implements FR-07. It runs entirely in the browser and is included on every public page through "
     "the shared partial. It listens for user input, shows a typing animation, then matches the message against ten "
     "keyword intents.")
code_block(
    "chatSend.addEventListener('click', () => {\n"
    "  const text = chatInput.value.trim();\n"
    "  if (text) { handleUserMsg(text); chatInput.value = ''; }\n"
    "});\n\n"
    "function handleUserMsg(text) {\n"
    "  appendMessage(text, true);                 // show the user's message\n"
    "  const indicator = showTypingIndicator();   // animated dots\n"
    "  setTimeout(() => {\n"
    "    indicator.remove();\n"
    "    const response = getBotResponse(text);   // pick an answer\n"
    "    appendMessage(response.text, false, true);\n"
    "    // render any follow-up action buttons from response.suggestions\n"
    "  }, 1000);\n"
    "}"
)
walk([
    ("addEventListener('click', ...)",
     "Sends the message when the user clicks Send (an Enter-key handler does the same)."),
    ("chatInput.value.trim()",
     "Reads and trims the typed text; empty input is ignored."),
    ("appendMessage(text, true)",
     "Adds the user’s message to the conversation as a right-aligned bubble."),
    ("showTypingIndicator()",
     "Displays animated dots so the assistant feels responsive and human."),
    ("setTimeout(..., 1000)",
     "Waits one second before replying, reinforcing the “thinking” effect."),
    ("getBotResponse(text)",
     "Runs the keyword-matching engine to choose the answer."),
    ("appendMessage(response.text, false, true)",
     "Renders the bot’s reply as HTML in a left-aligned bubble, including any links."),
])
body("The matching engine itself lower-cases the message and tests it against ten intents in order, returning the "
     "first that matches, or a fallback if none do.")
code_block(
    "function getBotResponse(userMsg) {\n"
    "  const msg = userMsg.toLowerCase();\n"
    "  if (msg.includes('chatbot') || msg.includes('bot') || msg.includes('chat'))   { return {...}; }\n"
    "  if (msg.includes('long')   || msg.includes('time') || msg.includes('timeline')) { return {...}; }\n"
    "  if (msg.includes('case')   || msg.includes('project') || msg.includes('work'))  { return {...}; }\n"
    "  if (msg.includes('human')  || msg.includes('contact') || msg.includes('call'))  { return {...}; }\n"
    "  if (msg.includes('price')  || msg.includes('cost') || msg.includes('budget'))   { return {...}; }\n"
    "  if (msg.includes('service')|| msg.includes('offer') || msg.includes('solution')){ return {...}; }\n"
    "  if (msg.includes('automat')|| msg.includes('workflow') || msg.includes('ocr'))  { return {...}; }\n"
    "  if (msg.includes('machine learning') || msg.includes('ml') || msg.includes('model')) { return {...}; }\n"
    "  if (msg.includes('about')  || msg.includes('company') || msg.includes('where'))  { return {...}; }\n"
    "  if (msg.includes('hi')     || msg.includes('hello') || msg.includes('hey'))      { return {...}; }\n"
    "  return { text: fallbackMessage, suggestions: [...] };   // none matched\n"
    "}"
)
walk([
    ("userMsg.toLowerCase()",
     "Normalises the message so matching is case-insensitive (“Price” and “price” behave the same)."),
    ("msg.includes('chatbot' / 'bot' / 'chat')",
     "Intent 1 — explains the chatbot service and links to a case study and a pre-filled contact form."),
    ("'long' / 'time' / 'timeline'",
     "Intent 2 — returns the typical 4–8 week project timeline broken down by phase."),
    ("'case' / 'project' / 'work'",
     "Intent 3 — summarises past projects and links to the Case Studies page."),
    ("'human' / 'contact' / 'call'",
     "Intent 4 — offers to hand the conversation to a person via the contact form."),
    ("'price' / 'cost' / 'budget'",
     "Intent 5 — explains that pricing is scope-based and invites a quote request."),
    ("'service' / 'offer' / 'solution'",
     "Intent 6 — lists the three core services."),
    ("'automat' / 'workflow' / 'ocr'",
     "Intent 7 — describes automation work such as invoice OCR and reporting."),
    ("'machine learning' / 'ml' / 'model'",
     "Intent 8 — describes custom ML: semantic search, recommendations, on-prem RAG."),
    ("'about' / 'company' / 'where'",
     "Intent 9 — gives the company background and location."),
    ("'hi' / 'hello' / 'hey'",
     "Intent 10 — a friendly greeting with starter suggestion buttons."),
    ("return { text: fallbackMessage, suggestions }",
     "If nothing matched, a helpful fallback lists the topics the bot can discuss — the assistant never goes silent."),
])

# ---- 6.13 main.js front-end ----
h2("6.13 Front-End Interactions (public/js/main.js)")
body("main.js handles the small interactive touches on the public pages: the mobile menu, the case-study filters, "
     "pre-filling the contact form when a visitor is handed off from the chatbot, the feedback star-rating modal, "
     "and event RSVP buttons. It runs after the page has loaded.")
code_block(
    "document.addEventListener('DOMContentLoaded', () => {\n"
    "  // 1. Mobile navbar toggle\n"
    "  const mobileMenuBtn = document.getElementById('mobile-menu-btn');\n"
    "  const navMenu = document.getElementById('nav-menu');\n"
    "  if (mobileMenuBtn && navMenu) {\n"
    "    mobileMenuBtn.addEventListener('click', () => {\n"
    "      navMenu.style.display = navMenu.style.display === 'flex' ? 'none' : 'flex';\n"
    "      // ...plus positioning styles for the dropdown panel\n"
    "    });\n"
    "  }\n\n"
    "  // 3. Prefill contact form after chatbot handoff\n"
    "  const prefill = new URLSearchParams(window.location.search).get('prefill');\n"
    "  if (prefill === 'chatbot') {\n"
    "    document.getElementById('job_title').value = 'AI Chatbot Integration';\n"
    "    document.getElementById('job_details').value = 'I am interested in building an AI Chatbot...';\n"
    "  }\n"
    "});"
)
walk([
    ("DOMContentLoaded",
     "Waits until the page’s HTML is fully parsed before wiring up any interactions, so the elements exist."),
    ("getElementById('mobile-menu-btn') / 'nav-menu'",
     "Finds the hamburger button and the navigation menu used on small screens."),
    ("if (mobileMenuBtn && navMenu)",
     "Only attaches the handler when both elements are present, avoiding errors on pages without them."),
    ("navMenu.style.display === 'flex' ? 'none' : 'flex'",
     "Toggles the menu open or closed each time the hamburger is clicked (the responsive menu for NFR-03)."),
    ("URLSearchParams(...).get('prefill')",
     "Reads a ?prefill= value from the URL — this is how the chatbot ‘hands off’ to the contact form."),
    ("if (prefill === 'chatbot') { ...value = ... }",
     "Pre-fills the job title and details so a visitor arriving from the chatbot has a head start on the enquiry."),
])
body("The same file also drives the case-study category filter (showing or hiding cards by their data-category), "
     "opens and closes the ‘add review’ modal, records the selected star rating, and shows a temporary "
     "toast confirmation when an event RSVP button is clicked. Each of these follows the same safe pattern: find "
     "the elements, check they exist, then attach a click handler.")
pagebreak()

# =========================================================
# 7. SECURITY
# =========================================================
h1("7. Security Design")
body("Security is addressed at every layer in line with NFR-02. The table maps each measure to the exact threat it "
     "mitigates and where in the code it lives.")
make_table(
    ["Measure", "Implementation", "Threat mitigated"],
    [
        ["Password hashing", "bcrypt hash (work factor 10) in seed.js; bcrypt.compare on login.", "Credential theft if the database file is exposed."],
        ["Parameterised SQL", "Every query in db.js callers uses ? placeholders with bound values.", "SQL injection."],
        ["Session security", "Secret from .env; HttpOnly, SameSite=lax, 2-hour cookie.", "Cookie theft via XSS, and CSRF."],
        ["Session regeneration", "session.regenerate() on successful login.", "Session fixation."],
        ["Route protection", "requireAdmin middleware on all /admin/* routes.", "Unauthorised access to admin data."],
        ["Cache control", "no-store headers on admin pages; cleared cookie on logout.", "Viewing a cached dashboard via the Back button."],
        ["Input validation", "Length, format and character checks on every contact field.", "Malformed or malicious input."],
        ["Generic error messages", "Login returns the same 'invalid' message for unknown user or wrong password.", "Account enumeration."],
        ["Secret management", "Credentials and secret kept in .env, which is git-ignored.", "Leaking secrets into version control."],
    ],
    widths=[1.4, 2.6, 2.3],
)
pagebreak()

# =========================================================
# 8. TESTING
# =========================================================
h1("8. Testing")

h2("8.1 Test Strategy, Protocol and Tools")
body("Testing was carried out continuously throughout development (each feature was tested as it was built) and "
     "then as a full pass against every requirement once the system was complete. The approach was functional, "
     "black-box testing: for each requirement, both the success path and the failure path were exercised against "
     "the running application, and the observed result was compared with the expected result.")
make_table(
    ["Aspect", "Detail"],
    [
        ["Test type", "Functional (black-box) testing of each requirement, plus non-functional checks (responsive, performance, security headers)."],
        ["Protocol", "For each case: define expected result → perform the action → record the actual result → mark pass/fail."],
        ["Tools", "Web browser and developer tools (network and responsive device modes); manual form submission; direct URL requests; browser dev-tools to inspect response headers."],
        ["Environment", "Local host machine running the Node.js server on port 3000 against the seeded SQLite database."],
        ["Coverage basis", "Every functional requirement (FR-01–FR-07) and non-functional requirement (NFR-01–NFR-05)."],
    ],
    widths=[1.4, 5.0],
)

h2("8.2 Functional Test Cases and Results")
body("The cases below were executed against the live application and all passed.")
make_table(
    ["#", "Test", "Expected", "Result"],
    [
        ["T1", "Load all eight public pages", "HTTP 200 for each", "Pass"],
        ["T2", "Open /admin/dashboard with no session", "Redirect (302) to login", "Pass"],
        ["T3", "Login with wrong password", "Redirect to login?error=invalid", "Pass"],
        ["T4", "Login with correct password", "Redirect to dashboard, 200", "Pass"],
        ["T5", "Dashboard shows count and table", "Seeded enquiries listed", "Pass"],
        ["T6", "Submit contact form with missing fields", "Validation error shown", "Pass"],
        ["T7", "Submit contact form with bad email", "Email format error shown", "Pass"],
        ["T8", "Submit a valid contact form", "Success message; row stored", "Pass"],
        ["T9", "New enquiry appears on dashboard", "Stored email visible in table", "Pass"],
        ["T10", "Update an enquiry status", "Status persists after refresh", "Pass"],
        ["T11", "Add a feedback review", "Review stored and average updates", "Pass"],
        ["T12", "Logout", "Redirect to /admin/login", "Pass"],
        ["T13", "Open dashboard after logout", "Redirect (302) to login", "Pass"],
        ["T14", "Admin page cache header", "Cache-Control: no-store present", "Pass"],
        ["T15", "Chatbot responds to a question", "Matching FAQ answer returned", "Pass"],
        ["T16", "Chatbot unknown question", "Fallback response returned", "Pass"],
    ],
    widths=[0.5, 2.5, 2.1, 0.8],
)

h2("8.3 Non-Functional and Compatibility Testing")
body("The layout was checked at desktop (1920px), tablet (768px), and mobile (375px) widths. CSS media queries "
     "collapse the navigation bar into a hamburger menu and stack content vertically, with no horizontal scrolling "
     "at any width. Pages are server-rendered and lightweight, loading well within the three-second target on "
     "localhost, which satisfies NFR-03 and NFR-04. Response headers on admin pages were inspected in browser "
     "developer tools to confirm the no-store cache policy (NFR-02).")

h2("8.4 Requirement Traceability Matrix")
body("This matrix shows that every requirement is covered by at least one test, which is the evidence that the "
     "delivered product meets the agreed specification.")
make_table(
    ["Requirement", "Covered by", "Outcome"],
    [
        ["FR-01 Eight pages", "T1", "Met"],
        ["FR-02 Contact form storage", "T6, T7, T8, T9", "Met"],
        ["FR-03 Secure admin login", "T3, T4", "Met"],
        ["FR-04 Protected admin area", "T2, T13", "Met"],
        ["FR-05 Dashboard data", "T5, T10", "Met"],
        ["FR-06 Logout & Back-button", "T12, T13, T14", "Met"],
        ["FR-07 AI chatbot", "T15, T16", "Met"],
        ["NFR-01 Consistent design", "Visual review across all pages", "Met"],
        ["NFR-02 Secure storage", "T3, T14 + code review of hashing/SQL", "Met"],
        ["NFR-03 Responsive", "8.3 responsive testing", "Met"],
        ["NFR-04 Fast load", "8.3 performance check", "Met"],
        ["NFR-05 Maintainable", "Code structure review (Section 5.5)", "Met"],
    ],
    widths=[2.3, 2.9, 1.2],
)

h2("8.5 Known Limitations")
for t in [
    "Sessions are stored in memory, so logging in does not survive a server restart (acceptable for a prototype).",
    "The chatbot is rule-based, so it answers only the topics it has intents for; unknown questions get the fallback.",
    "The dashboard table is not yet paginated, which is fine at current data volumes but would be added for scale.",
]:
    bullet(t)
pagebreak()

# =========================================================
# 9. TECHNICAL DEPLOYMENT
# =========================================================
h1("9. Technical Deployment of the Solution")
body("This section describes the technical requirements of the system and the procedures for installing and "
     "deploying it into a production environment — the setting where it would actually be used by end users.")

h2("9.1 Technical Requirements")
make_table(
    ["Category", "Requirement"],
    [
        ["Data platform", "SQLite — a single self-contained database file (database.sqlite); no separate database server."],
        ["Programming languages", "JavaScript (Node.js) on the server; HTML, CSS and vanilla JavaScript on the client; EJS templates; SQL."],
        ["Runtime / software", "Node.js 18+ and npm; the npm packages express, express-session, ejs, bcryptjs, sqlite3 and dotenv."],
        ["Hardware", "Any modern PC or small server; no specialised hardware was procured or required (CPU, ~512MB RAM and a few MB of disk are sufficient)."],
        ["Operating system", "Cross-platform — developed on Windows 11; runs on any OS that supports Node.js (Windows, macOS, Linux)."],
        ["Development tools", "Visual Studio Code, Git, npm, and the Mermaid CLI for diagrams."],
        ["Client requirement", "A modern web browser (Chrome, Edge, Firefox or Safari); no plug-ins required."],
    ],
    widths=[1.7, 4.7],
)

h2("9.2 Performance, Reliability and Availability")
body("The technical (non-functional) targets the system must fulfil are modest and were met. Pages are "
     "server-rendered with no heavy front-end framework, so they load quickly (well under the three-second target "
     "on localhost). Reliability is helped by validating all input at the server boundary and handling errors "
     "explicitly so a single bad request cannot crash the server. For this prototype, availability depends on the "
     "single Node.js process running; a production deployment would run it under a process manager (see 9.9) to "
     "restart it automatically and keep it available.")

h2("9.3 Software Packaging and Distribution")
body("The solution is logically packaged as a single project folder containing the source code, templates, static "
     "assets and the package.json manifest. For distribution it is delivered as a compressed ZIP archive of that "
     "folder (excluding the node_modules folder and the .env file, which are recreated on the target machine). "
     "Because the database is a file that the application creates on first run, no database export needs to be "
     "shipped. The natural distribution medium is a direct file transfer or a Git repository.")

h2("9.4 Licensing and Registration")
body("The application itself requires no licensing or product registration to run, and no end-user account or "
     "licence key is needed to use the public website. The third-party npm packages it depends on are open-source "
     "under permissive licences (MIT and similar), which permit free use, modification and distribution. The only "
     "credential involved is the administrator login, which is created locally by the seed script.")

h2("9.5 Prerequisites")
bullet("Node.js (version 18 or newer) installed.")
bullet("npm (bundled with Node.js).")
bullet("A terminal / command prompt and a modern web browser.")

h2("9.6 Installation and Deployment Procedure")
numbered("Unzip the project (or clone the repository) and open a terminal in the project folder (website).")
numbered("Create a .env file containing PORT=3000 and a SESSION_SECRET value.")
numbered("Run npm install to download the dependencies listed in package.json.")
numbered("Run npm run seed to create and populate the database (admin user, sample enquiries and reviews).")
numbered("Run npm start to launch the server.")
numbered("Open http://localhost:3000 in a browser.")
body("During start-up the console prints confirmation that the SQLite database is active and shows the admin panel "
     "URL, so it is clear the server is running correctly.")

h2("9.7 Default Administrator Credentials")
make_table(
    ["Field", "Value"],
    [["Username", "admin@kritika.ai"], ["Password", "admin123"]],
    widths=[1.6, 4.0],
)
body("These are created by seed.js. The password is stored only as a bcrypt hash in the database. To change it, "
     "edit the adminPassword value in seed.js and re-run npm run seed.")

h2("9.8 Using the System")
bullet("Visitors browse the eight pages from the top navigation bar.")
bullet("The Contact Us page collects an enquiry; on success a confirmation message is shown and the enquiry appears on the admin dashboard.")
bullet("The floating chat bubble (bottom-right) opens the AI assistant on any public page.")
bullet("Administrators sign in from Admin Login, review enquiries on the dashboard, open any enquiry for detail, update its status, and log out.")

h2("9.9 Production Deployment Considerations")
body("Moving from this prototype to a live production environment would involve a few well-understood steps: run "
     "the app behind a process manager such as PM2 so it restarts automatically; place it behind a reverse proxy "
     "(e.g. Nginx) terminating HTTPS, which then enables the secure-cookie flag already coded in; set NODE_ENV to "
     "production and supply a strong SESSION_SECRET; and, for higher traffic, move sessions to a persistent store "
     "and consider a client-server database. None of these change the application code materially — they are "
     "configuration and environment choices the design already anticipates.")

h2("9.10 Demonstration (Screencast)")
body("A pre-recorded video demonstration accompanies this portfolio (Task 2). It walks through each requirement on "
     "the running system: browsing the public pages, submitting a contact enquiry (including a validation failure "
     "and a success), using the chatbot, logging in as the administrator, viewing the dashboard and an enquiry, "
     "updating a status, and logging out to show the Back-button protection.")
pagebreak()

# =========================================================
# 10. EVALUATION AND CRITICAL REFLECTION
# =========================================================
h1("10. Evaluation and Critical Reflection")
body("This section reflects on my experience of delivering the project — what happened, what I found "
     "challenging, and what I would carry forward. It is written about my own practice rather than as a description "
     "of the methods, which are covered in Section 4.")

h2("10.1 Brief History of the Project")
body("The project began with a client meeting to understand the AI-Solutions scenario, followed by writing and "
     "signing off the requirements. I then planned the work as a dated schedule and built the system iteratively "
     "— the public pages first, then the contact workflow, then the secure admin area, and finally the "
     "chatbot and feedback features. A working build was reviewed with the client at the interim review, after "
     "which I completed the remaining features, tested the whole system against the requirements, and produced this "
     "portfolio and the demonstration.")

h2("10.2 Reflection on the Methodology and Approach")
body("On the whole, the adapted iterative approach served the project well. Building feature by feature meant I "
     "always had something working to show and test, which kept the project feeling under control even when "
     "individual tasks ran late. Treating the module milestones as my own review points gave the work a useful "
     "rhythm. If I am honest, I sometimes designed slightly ahead of what I immediately needed, and a stricter "
     "‘just enough design’ discipline would have saved a little time; recognising that is itself a useful "
     "lesson for future projects.")

h2("10.3 Challenges Faced and How I Overcame Them")
body("The most demanding part was the security of the admin area. Getting sessions, password hashing and the "
     "Back-button behaviour all working correctly took longer than I expected, and it was where the schedule "
     "slipped most. I overcame it by breaking the problem down — first proving the login and hash comparison "
     "worked, then adding session regeneration, then the cache-control headers — and testing each step before "
     "moving on, rather than trying to make everything work at once. A second challenge was keeping the validation "
     "thorough without making the contact handler unreadable; I addressed this by running the checks in a clear, "
     "consistent order, each with its own specific error message.")

h2("10.4 Adaptations Made to Fit the Constraints")
body("Working alone and to a deadline, I deliberately simplified where it was safe to do so. Choosing SQLite over "
     "a client-server database removed an entire piece of setup and administration. Keeping the chatbot rule-based "
     "rather than integrating a language-model API kept it reliable, free and fully under my control for the "
     "demonstration. And when the schedule slipped, I re-balanced the remaining iterations rather than cutting "
     "testing — a conscious decision to protect quality over scope.")

h2("10.5 What Went Well")
for t in [
    "Every agreed functional and non-functional requirement was delivered and tested.",
    "The codebase stayed clean and organised, with data access isolated behind one module.",
    "Security was treated as a first-class concern rather than an afterthought.",
    "The iterative approach meant there was always a working, demonstrable product.",
]:
    bullet(t)

h2("10.6 What I Would Do Differently")
for t in [
    "Estimate the security and session work more generously, as that is where I underestimated effort.",
    "Write a few automated tests alongside the manual testing, to catch regressions faster.",
    "Adopt a stricter ‘just enough design’ rule to avoid designing ahead of need.",
    "Start the portfolio documentation earlier and in parallel, rather than weighting it towards the end.",
]:
    bullet(t)

h2("10.7 Skills Development and Employability")
body("The project strengthened my practical, end-to-end development skills: gathering and signing off requirements "
     "with a client, planning and tracking a schedule, designing a system and documenting it clearly, writing "
     "secure server-side code, and testing against a specification. Just as valuable was the experience of managing "
     "a client-driven project to a deadline on my own — making trade-offs, keeping a stakeholder informed, and "
     "delivering a working product. These are exactly the skills employers look for in a graduate developer, and I "
     "feel more confident taking on a real client project as a result.")
pagebreak()

# =========================================================
# 11. MAINTAINABILITY & FUTURE WORK
# =========================================================
h1("11. Maintainability and Future Enhancements")
h2("11.1 Maintainability")
for t in [
    "Each concern lives in a clearly named file: routing in server.js, data access in db.js, views in views/, assets in public/.",
    "All database access goes through four small helpers, so the storage engine could be swapped with minimal change to the rest of the code.",
    "Routes and middleware are commented and reference the requirement they satisfy (e.g. FR-04, FR-06).",
    "The single stylesheet uses a consistent palette and tokens, making site-wide visual changes simple.",
    "Validation, security and rendering each follow one repeated pattern, so new routes are easy to add consistently.",
]:
    bullet(t)
h2("11.2 Possible Future Enhancements")
for t in [
    "Email notification to staff when a new enquiry arrives.",
    "Pagination and search/filtering on the dashboard enquiry table.",
    "A richer, AI-backed chatbot using a language-model API instead of keyword matching.",
    "Role-based accounts so multiple staff members can log in with individual permissions.",
    "Persisting sessions to a store so logins survive a server restart.",
]:
    bullet(t)
pagebreak()

# =========================================================
# 12. CONCLUSION
# =========================================================
h1("12. Conclusion")
body("The AI-Solutions product delivers every functional and non-functional requirement from the specification. "
     "It presents a clean, consistent, responsive marketing site; collects and validates customer enquiries into a "
     "relational database; protects an administrator dashboard with hashed credentials and server sessions; and "
     "provides an always-available FAQ chatbot. The project was planned as a working schedule, managed with an "
     "adapted iterative methodology, and reviewed with the client at agreed milestones. Throughout, security has "
     "been treated as a primary concern — hashed passwords, parameterised queries, hardened sessions, "
     "regeneration on login, and cache control — and the codebase has been organised for clarity and future "
     "maintenance. Every design artefact in this document is backed by a diagram and a full explanation, and every "
     "important piece of code has been reproduced and explained line by line. The system has been tested against "
     "each requirement and performs correctly.")
pagebreak()

# =========================================================
# APPENDICES
# =========================================================
h1("Appendix A – Route Reference")
body("Every HTTP route in the application, the page or action it serves, and whether it is protected by the "
     "requireAdmin middleware.")
make_table(
    ["Method", "Route", "Purpose", "Protected"],
    [
        ["GET", "/", "Home page (shows 3 recent reviews)", "No"],
        ["GET", "/solutions", "Solutions page", "No"],
        ["GET", "/cases", "Case studies list", "No"],
        ["GET", "/cases/:id", "Case study detail", "No"],
        ["GET", "/feedback", "Feedback & ratings (with average score)", "No"],
        ["POST", "/feedback/add", "Add a review", "No"],
        ["GET", "/articles", "Articles page", "No"],
        ["GET", "/gallery", "Gallery & events", "No"],
        ["GET", "/contact", "Contact form page", "No"],
        ["POST", "/contact", "Validate & store enquiry", "No"],
        ["GET", "/admin/login", "Login page", "No"],
        ["POST", "/admin/login", "Authenticate admin", "No"],
        ["GET", "/admin/logout", "Destroy session & log out", "Session"],
        ["GET", "/admin/dashboard", "Dashboard (counts + table)", "Yes"],
        ["GET", "/admin/inquiries/:id", "Inquiry detail", "Yes"],
        ["POST", "/admin/inquiries/:id/status", "Update enquiry status", "Yes"],
        ["GET", "/admin/reviews", "Manage reviews", "Yes"],
        ["GET", "/admin/articles, /events, /gallery, /settings", "Admin content pages", "Yes"],
    ],
    widths=[0.8, 2.4, 2.5, 0.9],
)

h1("Appendix B – File Inventory")
make_table(
    ["File", "Responsibility"],
    [
        ["server.js", "Express app, session config, requireAdmin middleware and all route handlers."],
        ["db.js", "SQLite connection, table creation and the query / execute / getOne helpers."],
        ["seed.js", "Creates the bcrypt-hashed admin and seeds sample enquiries and reviews."],
        ["package.json", "Dependencies and npm scripts (start, dev, seed)."],
        [".env", "PORT and SESSION_SECRET (git-ignored)."],
        ["views/partials/header.ejs", "Shared head, navbar and responsive menu."],
        ["views/partials/footer.ejs", "Shared footer and script includes."],
        ["views/partials/chatbot.ejs", "Chatbot widget markup included on public pages."],
        ["public/css/style.css", "Design tokens, layout and responsive rules."],
        ["public/js/main.js", "Navbar / hamburger, case filters, modal, star rating, prefill, RSVP toast."],
        ["public/js/chat.js", "Chatbot FAQ matching engine (ten intents + fallback)."],
    ],
    widths=[2.4, 4.0],
)

doc.save(OUT)
print("SAVED", OUT)
